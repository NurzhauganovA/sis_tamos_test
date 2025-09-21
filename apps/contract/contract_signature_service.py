# apps/contract/contract_signature_service.py

import hashlib
import base64
import os
import subprocess
import json
from datetime import datetime
from io import BytesIO

import qrcode
import requests
from typing import Dict, Any, Optional
import logging

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from docx import Document
from docx.shared import Inches, Cm

from .models import ContractSignature, ContractMS, ContractFileUser, ContractStatusMS, ContractDopMS, \
    ContractDopFileUser
from django.contrib.auth.models import User

logger = logging.getLogger(__name__)


class ContractSignatureService:
    """Обновленный сервис для работы с подписями контрактов"""

    def __init__(self):
        # URL FastAPI сервиса для верификации подписей
        self.fastapi_verify_url = getattr(
            settings,
            'FASTAPI_SIGNATURE_VERIFY_URL',
            'https://cabinet.tamos-education.kz:11443/fastapi/api/v1/contracts/verify-signature'
        )
        self.request_timeout = 30.0
        self.frontend_url = getattr(settings, 'FRONTEND_URL', 'https://cabinet.tamos-education.kz:11443')

    def verify_and_save_signature(
            self,
            contract_num: str,
            cms_signature: str,
            signed_data: str,
            user: User,
            is_dop_contract: bool = False
    ) -> Dict[str, Any]:
        """
        Верифицирует подпись через FastAPI и сохраняет в базу с обновлением PDF
        """
        try:
            # Находим контракт по номеру
            try:
                if is_dop_contract:
                    # Для дополнительных договоров
                    contract_dop = ContractDopMS.objects.using('ms_sql').filter(
                        agreement_id__ContractNum=contract_num
                    ).first()
                    if not contract_dop:
                        return {
                            'success': False,
                            'error': 'Дополнительный договор не найден',
                            'error_code': 'CONTRACT_NOT_FOUND'
                        }
                    contract = contract_dop.agreement_id
                else:
                    contract = ContractMS.objects.using('ms_sql').get(ContractNum=contract_num)
            except ContractMS.DoesNotExist:
                return {
                    'success': False,
                    'error': 'Контракт не найден',
                    'error_code': 'CONTRACT_NOT_FOUND'
                }

            # Проверяем, нет ли уже валидной подписи для этого контракта
            existing_signature = ContractSignature.objects.filter(
                contract_num=contract_num,
                is_valid=True
            ).first()

            if existing_signature and not existing_signature.is_document_modified:
                return {
                    'success': False,
                    'error': 'Контракт уже подписан',
                    'error_code': 'ALREADY_SIGNED'
                }

            # Верифицируем подпись через FastAPI
            verification_result = self._verify_signature_via_fastapi(
                cms_signature, signed_data
            )

            if not verification_result['success']:
                return verification_result

            # Проверяем соответствие ИИН
            if hasattr(user, 'user_info') and hasattr(user.user_info, 'iin'):
                if verification_result['iin'] != user.user_info.iin:
                    return {
                        'success': False,
                        'error': 'ИИН подписанта не совпадает с ИИН пользователя',
                        'error_code': 'IIN_MISMATCH'
                    }

            logger.info(
                f"Signature verification successful for contract {contract_num}, IIN: {verification_result['iin']}")

            # АТОМАРНАЯ ТРАНЗАКЦИЯ: Создаем файл и подпись как единое целое
            with transaction.atomic():
                try:
                    # 1. Генерируем полный подписанный контракт СНАЧАЛА
                    self._generate_complete_signed_contract_for_signature(
                        contract=contract,
                        user=user,
                        is_dop_contract=is_dop_contract,
                        signer_iin=verification_result['iin']
                    )

                    logger.info(f"Contract PDF generated successfully for {contract_num}")

                    # 2. ПОСЛЕ создания файла вычисляем хэш НОВОГО документа
                    document_hash = self._calculate_contract_hash(contract, is_dop_contract)

                    logger.info(f"New document hash calculated: {document_hash[:16]}...")

                    # 3. ТОЛЬКО ТЕПЕРЬ сохраняем подпись с правильным хэшем
                    signature = ContractSignature.objects.create(
                        contract_num=contract_num,
                        cms_signature=cms_signature,
                        signed_data=signed_data,
                        document_hash=document_hash,
                        signer_iin=verification_result['iin'],
                        certificate_info=verification_result.get('certificate_info', {}),
                        is_valid=True,
                        created_by=user
                    )

                    logger.info(f"ContractSignature created with hash: {document_hash[:16]}...")

                    # 4. Обновляем статус контракта
                    if is_dop_contract:
                        contract_dop.status_id = ContractStatusMS.objects.using('ms_sql').get(sStatusName='Подписан')
                        contract_dop.save(using='ms_sql')
                    else:
                        contract.ContractStatusID = ContractStatusMS.objects.using('ms_sql').get(sStatusName='Подписан')
                        contract.save(using='ms_sql')

                    # 5. Добавляем подпись директора (автоматически) с тем же хэшем
                    self._add_director_signature(contract_num, signature, document_hash, signed_data=signed_data)

                    logger.info(f"Transaction completed successfully for contract {contract_num}")

                except Exception as e:
                    logger.error(f"Error in transaction for contract {contract_num}: {e}")
                    # Транзакция автоматически откатится
                    raise

            logger.info(f"Signature saved successfully for contract {contract_num}, IIN: {verification_result['iin']}")

            return {
                'success': True,
                'signature_uid': str(signature.signature_uid),
                'signer_iin': verification_result['iin'],
                'contract_num': contract.ContractNum,
                'message': 'Подпись успешно верифицирована и сохранена'
            }

        except Exception as e:
            logger.error(f"Error in verify_and_save_signature: {e}")
            return {
                'success': False,
                'error': f'Ошибка при обработке подписи: {str(e)}',
                'error_code': 'PROCESSING_ERROR'
            }

    def _generate_complete_signed_contract_for_signature(self, contract, user, is_dop_contract=False, signer_iin=None):
        """Генерирует полный подписанный контракт специально для процесса подписания"""
        try:
            # Получаем данные студента и родителя
            student = contract.StudentID
            parent = student.parent_id if student else None

            if not student or not parent:
                logger.error("Student or parent not found for contract generation")
                raise ValueError("Student or parent data missing")

            # Генерируем временные QR-коды для подписи (будут обновлены после сохранения подписи)
            temp_qr_data = {
                "type": "contract_signature",
                "contract_num": contract.ContractNum,
                "signer_iin": signer_iin or "pending",
                "signed_at": datetime.now().isoformat(),
                "message": "Подпись в процессе обработки"
            }
            qr_signature_code = self._create_qr_code(temp_qr_data)

            # Генерируем QR-коды директоров
            qr_director_omarov = self._generate_director_qr_code('omarov', contract.ContractNum)
            qr_director_serikov = self._generate_director_qr_code('serikov', contract.ContractNum)

            # Создаем новый документ с заполненными переменными и QR-кодами
            self._generate_complete_signed_contract(
                contract=contract,
                student=student,
                parent=parent,
                qr_signature=qr_signature_code,
                qr_director_omarov=qr_director_omarov,
                qr_director_serikov=qr_director_serikov,
                user=user,
                is_dop_contract=is_dop_contract
            )

            logger.info(f"Contract file generated successfully for signing process: {contract.ContractNum}")

        except Exception as e:
            logger.error(f"Error in _generate_complete_signed_contract_for_signature: {e}")
            raise

    def _verify_signature_via_fastapi(
            self,
            cms_signature: str,
            signed_data: str
    ) -> Dict[str, Any]:
        """Верифицирует подпись через FastAPI сервис"""
        try:
            payload = {
                'cms': cms_signature,
                'data': signed_data
            }

            logger.info(f"Sending verification request to FastAPI: {self.fastapi_verify_url}")

            response = requests.post(
                self.fastapi_verify_url,
                json=payload,
                timeout=self.request_timeout,
                headers={'Content-Type': 'application/json'}
            )

            logger.info(f"FastAPI response status: {response.status_code}")

            if response.status_code == 200:
                result = response.json()

                if result.get('success'):
                    return {
                        'success': True,
                        'iin': result['iin'],
                        'certificate_info': result.get('certificate_info', {})
                    }
                else:
                    return {
                        'success': False,
                        'error': result.get('error', 'Ошибка верификации подписи'),
                        'error_code': result.get('error_code', 'VERIFICATION_FAILED')
                    }
            else:
                try:
                    error_data = response.json()
                    error_message = error_data.get('detail', {})
                    if isinstance(error_message, dict):
                        error_message = error_message.get('error', 'Ошибка верификации подписи')
                except:
                    error_message = f'HTTP {response.status_code}: {response.text}'

                return {
                    'success': False,
                    'error': error_message,
                    'error_code': 'VERIFICATION_FAILED'
                }

        except requests.exceptions.Timeout:
            logger.error("Timeout while verifying signature")
            return {
                'success': False,
                'error': 'Таймаут при верификации подписи',
                'error_code': 'TIMEOUT'
            }
        except requests.exceptions.ConnectionError as e:
            logger.error(f"Connection error while verifying signature: {e}")
            return {
                'success': False,
                'error': 'Ошибка соединения с сервисом верификации',
                'error_code': 'CONNECTION_ERROR'
            }
        except Exception as e:
            logger.error(f"Unexpected error while verifying signature: {e}")
            return {
                'success': False,
                'error': f'Неожиданная ошибка при верификации: {str(e)}',
                'error_code': 'UNEXPECTED_ERROR'
            }

    def _update_contract_pdf_with_signature(self, contract, signature, user, is_dop_contract=False):
        """Создает полностью новый PDF контракта с заполненными переменными и QR-кодами"""
        try:
            # Всегда генерируем новый контракт с заполненными данными
            logger.info(f"Generating new contract file for {contract.ContractNum}")

            # Получаем данные студента и родителя
            student = contract.StudentID
            parent = student.parent_id if student else None

            if not student or not parent:
                logger.error("Student or parent not found for contract generation")
                return

            # Генерируем QR-коды
            qr_signature_data = self._generate_signature_qr_data(signature)
            qr_signature_code = self._create_qr_code(qr_signature_data)
            qr_director_omarov = self._generate_director_qr_code('omarov', contract.ContractNum)
            qr_director_serikov = self._generate_director_qr_code('serikov', contract.ContractNum)

            # Создаем новый документ с заполненными переменными и QR-кодами
            self._generate_complete_signed_contract(
                contract=contract,
                student=student,
                parent=parent,
                qr_signature=qr_signature_code,
                qr_director_omarov=qr_director_omarov,
                qr_director_serikov=qr_director_serikov,
                user=user,
                is_dop_contract=is_dop_contract
            )

            logger.info(f"Contract PDF created successfully for {contract.ContractNum}")

        except Exception as e:
            logger.error(f"Error creating contract PDF: {e}")
            # Не прерываем процесс подписания из-за ошибки обновления PDF

    def _generate_complete_signed_contract(self, contract, student, parent, qr_signature, qr_director_omarov,
                                           qr_director_serikov, user, is_dop_contract=False):
        """Генерирует полный подписанный контракт с заполненными переменными и QR-кодами"""
        try:
            # Получаем правильный шаблон контракта
            docx_template_path = self._get_contract_template(contract, is_dop_contract)

            if not docx_template_path:
                logger.error("Contract template not found")
                return

            # Проверяем существование файла шаблона
            import os
            if not os.path.exists(docx_template_path):
                logger.error(f"Template file not found: {docx_template_path}")
                return

            # Открываем шаблон
            doc = Document(docx_template_path)

            # Заполняем ВСЕ переменные контракта И добавляем QR-коды
            self._replace_qr_placeholders(doc, qr_signature, qr_director_omarov, qr_director_serikov, contract)

            # Обрабатываем специальные таблицы оплаты если есть
            self._process_payment_tables(doc, contract)

            # Сохраняем готовый документ
            docx_output_path = f'contracts/signed/docx/contract_{contract.ContractNum}_signed.docx'
            pdf_directory = "contracts/signed/pdf"
            pdf_output_path = f'{pdf_directory}/contract_{contract.ContractNum}_signed.pdf'

            # Создаем директории если не существуют
            os.makedirs(os.path.dirname(docx_output_path), exist_ok=True)
            os.makedirs(pdf_directory, exist_ok=True)

            # Сохраняем DOCX
            doc.save(docx_output_path)
            logger.info(f"DOCX saved: {docx_output_path}")

            # Конвертируем в PDF
            self._docx_to_pdf(docx_output_path, pdf_directory)
            logger.info(f"PDF converted: {pdf_output_path}")

            # Проверяем что PDF создался
            if not os.path.exists(pdf_output_path):
                logger.error(f"PDF file was not created: {pdf_output_path}")
                return

            # Сохраняем в базе данных (создаем новую запись или обновляем существующую)
            with open(pdf_output_path, 'rb') as pdf_file:
                file_content = pdf_file.read()

                if is_dop_contract:
                    # Для дополнительных договоров
                    contract_file = ContractDopFileUser.objects.filter(contractNum=contract.ContractNum).first()
                    if contract_file:
                        # Обновляем существующую запись
                        contract_file.file.delete(save=False)  # Удаляем старый файл
                        contract_file.file = ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        contract_file.date = datetime.now()
                        contract_file.save()
                        logger.info(f"Updated existing dop contract file: {contract.ContractNum}")
                    else:
                        # Создаем новую запись
                        ContractDopFileUser.objects.create(
                            user=user,
                            contractNum=contract.ContractNum,
                            file=ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        )
                        logger.info(f"Created new dop contract file: {contract.ContractNum}")
                else:
                    # Для основных договоров
                    contract_file = ContractFileUser.objects.filter(contractNum=contract.ContractNum).first()
                    if contract_file:
                        # Обновляем существующую запись
                        contract_file.file.delete(save=False)  # Удаляем старый файл
                        contract_file.file = ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        contract_file.date = datetime.now()
                        contract_file.save()
                        logger.info(f"Updated existing contract file: {contract.ContractNum}")
                    else:
                        # Создаем новую запись
                        ContractFileUser.objects.create(
                            user=user,
                            contractNum=contract.ContractNum,
                            file=ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        )
                        logger.info(f"Created new contract file: {contract.ContractNum}")

            # Удаляем временные файлы
            try:
                os.remove(docx_output_path)
                os.remove(pdf_output_path)
                logger.info("Temporary files cleaned up")
            except Exception as e:
                logger.warning(f"Could not remove temporary files: {e}")

        except Exception as e:
            logger.error(f"Error generating complete signed contract: {e}")
            raise

    def _process_payment_tables(self, doc, contract):
        """Обрабатывает специальные таблицы оплаты в документе"""
        try:
            # Ищем и обрабатываем таблицы в документе
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if '{customtable_monthpay}' in run.text:
                                    self._process_month_pay_table(cell, contract)
                                    run.text = run.text.replace('{customtable_monthpay}', '')

                                if '{customtable_quarterpay}' in run.text:
                                    self._process_quarter_pay_table(cell, contract)
                                    run.text = run.text.replace('{customtable_quarterpay}', '')

        except Exception as e:
            logger.error(f"Error processing payment tables: {e}")

    def _generate_signature_qr_data(self, signature):
        """Генерирует данные для QR-кода подписи"""
        # URL для проверки подписи на фронтенде
        verification_url = f"{self.frontend_url}/signature-verification/{signature.signature_uid}"

        qr_data = {
            "type": "contract_signature",
            "signature_uid": str(signature.signature_uid),
            "contract_num": signature.contract_num,
            "signer_iin": signature.signer_iin,
            "signed_at": signature.signed_at.isoformat(),
            "verification_url": verification_url,
            "message": "Сканируйте для проверки подписи контракта"
        }

        return qr_data

    def _create_qr_code(self, data):
        """Создает QR-код из данных"""
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4
        )

        # Преобразуем данные в JSON строку для QR-кода
        qr_text = json.dumps(data, ensure_ascii=False)

        qr.add_data(qr_text)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")

        buffered = BytesIO()
        img.save(buffered, format='PNG')

        return buffered.getvalue()

    def _generate_director_qr_code(self, director_type, contract_num):
        """Генерирует QR-код директора"""
        try:
            # Данные директора для QR-кода
            if director_type == 'omarov':
                director_data = {
                    "type": "director_signature",
                    "director": "ОМАРОВ",
                    "position": "Директор",
                    "contract_num": contract_num,
                    "signed_at": datetime.now().isoformat(),
                    "certificate_info": {
                        "serial_number": "IIN540217301387",
                        "common_name": "ОМАРОВ МУРАТ"
                    }
                }
            elif director_type == 'serikov':
                director_data = {
                    "type": "director_signature",
                    "director": "СЕРИКОВ",
                    "position": "Директор",
                    "contract_num": contract_num,
                    "signed_at": datetime.now().isoformat(),
                    "certificate_info": {
                        "serial_number": "IIN861205300997",
                        "common_name": "СЕРИКОВ БАУЫРЖАН"
                    }
                }
            else:
                return b''

            return self._create_qr_code(director_data)

        except Exception as e:
            logger.error(f"Error generating director QR code: {e}")
            return b''

    def _add_qr_codes_to_contract(self, contract, qr_signature, qr_director_omarov, qr_director_serikov, user,
                                  is_dop_contract=False):
        """Добавляет QR-коды в документ контракта"""
        try:
            # Получаем шаблон контракта
            docx_template = self._get_contract_template(contract, is_dop_contract)

            if not docx_template:
                logger.error("Contract template not found")
                return

            doc = Document(docx_template)

            # Заменяем плейсхолдеры QR-кодов на реальные изображения
            self._replace_qr_placeholders(doc, qr_signature, qr_director_omarov, qr_director_serikov)

            # Сохраняем обновленный документ
            docx_output_path = f'contracts/signed/docx/contract_{contract.ContractNum}.docx'
            pdf_directory = "contracts/signed/pdf"
            pdf_output_path = f'{pdf_directory}/contract_{contract.ContractNum}.pdf'

            # Создаем директории если не существуют
            os.makedirs(os.path.dirname(docx_output_path), exist_ok=True)
            os.makedirs(pdf_directory, exist_ok=True)

            doc.save(docx_output_path)

            # Конвертируем в PDF
            self._docx_to_pdf(docx_output_path, pdf_directory)

            # Обновляем файл в базе данных
            with open(pdf_output_path, 'rb') as pdf_file:
                file_content = pdf_file.read()

                if is_dop_contract:
                    contract_file = ContractDopFileUser.objects.filter(contractNum=contract.ContractNum).last()
                    if contract_file:
                        contract_file.file = ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        contract_file.date = datetime.now()
                        contract_file.save()
                    else:
                        ContractDopFileUser.objects.create(
                            user=user,
                            contractNum=contract.ContractNum,
                            file=ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        )
                else:
                    contract_file = ContractFileUser.objects.filter(contractNum=contract.ContractNum).last()
                    if contract_file:
                        contract_file.file = ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        contract_file.date = datetime.now()
                        contract_file.save()
                    else:
                        ContractFileUser.objects.create(
                            user=user,
                            contractNum=contract.ContractNum,
                            file=ContentFile(file_content, name=f'{contract.ContractNum}_signed.pdf')
                        )

            # Удаляем временные файлы
            os.remove(docx_output_path)
            os.remove(pdf_output_path)

        except Exception as e:
            logger.error(f"Error adding QR codes to contract: {e}")

    def _replace_qr_placeholders(self, doc, qr_signature, qr_director_omarov, qr_director_serikov, contract):
        """Заменяет все плейсхолдеры в документе: переменные контракта и QR-коды"""
        from .models import ContractMS, ContractDopMS, DiscountMS, ContractDiscountMS
        from .services import ChangeDocumentContentService
        from datetime import timedelta
        import math
        from num2words import num2words
        from io import BytesIO

        # Получаем данные контракта
        try:
            if not contract:
                logger.error("Contract not provided for template replacement")
                return

            student = contract.StudentID
            parent = student.parent_id if student else None

            if not student or not parent:
                logger.error("Student or parent not found for contract")
                return

        except Exception as e:
            logger.error(f"Error getting contract data: {e}")
            return

        # Подготавливаем данные для замены
        try:
            # Данные паспорта родителя
            ParentPassport = f'Удостоверение личности: №{parent.num_of_doc}, Орган выдачи: {parent.issued_by}, Дата выдачи: {parent.issue_date}'
            ParentPassportKAZ = f'Жеке куәлік: №{parent.num_of_doc}, Берген орган: {parent.issued_by}, Берілген күні: {parent.issue_date}'
            ParentPassportENG = f'ID: No.{parent.num_of_doc}, Issued by: {parent.issued_by}, Issue date: {parent.issue_date}'

            if '?' in str(parent.issued_by):
                ParentPassport = f'Удостоверение личности: №{parent.num_of_doc}, Орган выдачи: Не указано, Дата выдачи: {parent.issue_date}'
                ParentPassportKAZ = f'Жеке куәлік: №{parent.num_of_doc}, Берген орган: Көрсетілмеген, Берілген күні: {parent.issue_date}'
                ParentPassportENG = f'ID: No.{parent.num_of_doc}, Issued by: Not specified, Issue date: {parent.issue_date}'

            # Расчет суммы со скидкой
            discount = DiscountMS.objects.using('ms_sql').all()
            contract_discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract)
            contract_amount = float(contract.ContractAmount) if contract.ContractAmount else 0

            if contract_discount.exists():
                for i in contract_discount:
                    for j in discount.filter(id=i.DiscountID.id):
                        percent = int(j.iDiscountPercent)
                        contract_amount -= contract_amount * percent / 100
                contract_amount_with_discount = float(round(contract_amount, 2))
            else:
                contract_amount_with_discount = float(round(contract_amount, 2))

            # Дополнительная сумма договора
            try:
                contract_dop_amount = ContractMS.objects.using('ms_sql').filter(
                    ContractNum=contract.ContractNum).first()
                if contract_dop_amount:
                    contract_dop_amount = contract_dop_amount.ContractAmount
                else:
                    contract_dop_amount = 0
            except:
                contract_dop_amount = 0

        except Exception as e:
            logger.error(f"Error calculating contract amounts: {e}")
            contract_amount_with_discount = float(contract.ContractAmount) if contract.ContractAmount else 0
            contract_dop_amount = 0
            ParentPassport = "Данные паспорта недоступны"
            ParentPassportKAZ = "Паспорт деректері қолжетімсіз"
            ParentPassportENG = "Passport data unavailable"

        # Заменяем переменные в обычных параграфах (как в change_docx_document)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                self._replace_text_variables(run, contract, student, parent, ParentPassport,
                                             ParentPassportKAZ, ParentPassportENG, contract_amount_with_discount,
                                             contract_dop_amount)
                self._replace_qr_codes(run, qr_signature, qr_director_omarov, qr_director_serikov, contract)

        # Заменяем переменные в таблицах (как в change_docx_document)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._replace_text_variables(run, contract, student, parent, ParentPassport,
                                                         ParentPassportKAZ, ParentPassportENG,
                                                         contract_amount_with_discount,
                                                         contract_dop_amount)
                            self._replace_qr_codes(run, qr_signature, qr_director_omarov, qr_director_serikov, contract)

                            # Обработка специальных таблиц
                            if '{customtable_monthpay}' in run.text:
                                self._process_month_pay_table(cell, contract)
                                run.text = run.text.replace('{customtable_monthpay}', '')

                            if '{customtable_quarterpay}' in run.text:
                                self._process_quarter_pay_table(cell, contract)
                                run.text = run.text.replace('{customtable_quarterpay}', '')

    def _replace_text_variables(self, run, contract, student, parent, ParentPassport, ParentPassportKAZ,
                                ParentPassportENG, contract_amount_with_discount, contract_dop_amount):
        """Заменяет все текстовые переменные в run (точно как в change_docx_document)"""
        from .services import ChangeDocumentContentService
        from datetime import timedelta
        from num2words import num2words

        try:
            # Основные данные контракта
            if '{ContractNum}' in run.text:
                run.text = run.text.replace('{ContractNum}', str(contract.ContractNum) if contract.ContractNum else '')
            if '{ContractYear}' in run.text:
                run.text = run.text.replace('{ContractYear}',
                                            str(contract.ContractDate.strftime("%Y")) if contract.ContractDate else '')
            if '{ContractYearFinish}' in run.text and contract.ContractDate:
                data_close = contract.ContractDate + timedelta(days=365)
                run.text = run.text.replace('{ContractYearFinish}', str(data_close.strftime("%Y")))
            if '{ContractDate}' in run.text:
                run.text = run.text.replace('{ContractDate}', str(contract.ContractDate.strftime(
                    "%d.%m.%Y")) if contract.ContractDate else '')
            if '{ContractDay}' in run.text:
                run.text = run.text.replace('{ContractDay}',
                                            str(contract.ContractDate.strftime("%d")) if contract.ContractDate else '')

            # Месяцы на разных языках
            if contract.ContractDate:
                if '{ContractMonthRUS}' in run.text:
                    try:
                        month_ru = ChangeDocumentContentService.translate_text(contract.ContractDate.strftime("%B"),
                                                                               "ru")
                        run.text = run.text.replace('{ContractMonthRUS}', month_ru)
                    except:
                        run.text = run.text.replace('{ContractMonthRUS}', contract.ContractDate.strftime("%B"))
                if '{ContractMonthKAZ}' in run.text:
                    try:
                        month_kz = ChangeDocumentContentService.translate_text(contract.ContractDate.strftime("%B"),
                                                                               "kk")
                        run.text = run.text.replace('{ContractMonthKAZ}', month_kz)
                    except:
                        run.text = run.text.replace('{ContractMonthKAZ}', contract.ContractDate.strftime("%B"))
                if '{ContractMonthENG}' in run.text:
                    try:
                        month_en = ChangeDocumentContentService.translate_text(contract.ContractDate.strftime("%B"),
                                                                               "en")
                        run.text = run.text.replace('{ContractMonthENG}', month_en)
                    except:
                        run.text = run.text.replace('{ContractMonthENG}', contract.ContractDate.strftime("%B"))

            # Учебный год
            if '{EduYear}' in run.text and contract.EduYearID:
                run.text = run.text.replace('{EduYear}', str(contract.EduYearID.sEduYear))

            # Данные студента и родителя
            if '{ParentFullName}' in run.text:
                run.text = run.text.replace('{ParentFullName}', str(parent.full_name) if parent.full_name else '')
            if '{StudentFullName}' in run.text:
                run.text = run.text.replace('{StudentFullName}', str(student.full_name) if student.full_name else '')
            if '{StudentIIN}' in run.text:
                run.text = run.text.replace('{StudentIIN}', str(student.iin) if student.iin else '')
            if '{StudentAddress}' in run.text:
                run.text = run.text.replace('{StudentAddress}', str(parent.address) if parent.address else '')
            if '{StudentPhoneNumber}' in run.text:
                phone_text = str(student.phone) if student.phone else '-'
                run.text = run.text.replace('{StudentPhoneNumber}', phone_text)
            if '{ParentAddress}' in run.text:
                run.text = run.text.replace('{ParentAddress}', str(parent.address) if parent.address else '')
            if '{ParentPhoneNumber}' in run.text:
                run.text = run.text.replace('{ParentPhoneNumber}', str(parent.phone) if parent.phone else '')
            if '{ParentIIN}' in run.text:
                run.text = run.text.replace('{ParentIIN}', str(parent.iin) if parent.iin else '')

            # Данные паспорта
            if '{ParentPassport}' in run.text:
                run.text = run.text.replace('{ParentPassport}', ParentPassport)
            if '{ParentPassportKAZ}' in run.text:
                run.text = run.text.replace('{ParentPassportKAZ}', ParentPassportKAZ)
            if '{ParentPassportENG}' in run.text:
                run.text = run.text.replace('{ParentPassportENG}', ParentPassportENG)

            # Суммы контракта
            if '{ContractAmount}' in run.text:
                run.text = run.text.replace('{ContractAmount}',
                                            str(int(contract.ContractAmount)) if contract.ContractAmount else '0')
            if '{ContractSum}' in run.text:
                run.text = run.text.replace('{ContractSum}',
                                            str(int(contract.ContractSum)) if contract.ContractSum else '0')
            if '{ContractAmountWithDiscount}' in run.text:
                run.text = run.text.replace('{ContractAmountWithDiscount}', str(int(contract_amount_with_discount)))

            # Дополнительные суммы
            if '{ContractDopAmount}' in run.text:
                run.text = run.text.replace('{ContractDopAmount}', str(int(contract_dop_amount)))

            # Взнос
            if '{ContractContr}' in run.text:
                contr_sum = int(contract.ContSum) if contract.ContSum else 0
                run.text = run.text.replace('{ContractContr}', str(contr_sum))

            # Суммы прописью
            try:
                if '{ContractAmountWords}' in run.text:
                    run.text = run.text.replace('{ContractAmountWords}',
                                                num2words(int(contract.ContractAmount),
                                                          lang="ru") if contract.ContractAmount else '')
                if '{ContractAmountWordsKaz}' in run.text:
                    run.text = run.text.replace('{ContractAmountWordsKaz}',
                                                num2words(int(contract.ContractAmount),
                                                          lang="kz") if contract.ContractAmount else '')
                if '{ContractAmountWordsEng}' in run.text:
                    run.text = run.text.replace('{ContractAmountWordsEng}',
                                                num2words(int(contract.ContractAmount),
                                                          lang="en") if contract.ContractAmount else '')
                if '{ContractSumWords}' in run.text:
                    run.text = run.text.replace('{ContractSumWords}',
                                                num2words(int(contract.ContractSum),
                                                          lang="ru") if contract.ContractSum else '')
                if '{ContractSumWordsKaz}' in run.text:
                    run.text = run.text.replace('{ContractSumWordsKaz}',
                                                num2words(int(contract.ContractSum),
                                                          lang="kz") if contract.ContractSum else '')
                if '{ContractSumWordsEng}' in run.text:
                    run.text = run.text.replace('{ContractSumWordsEng}',
                                                num2words(int(contract.ContractSum),
                                                          lang="en") if contract.ContractSum else '')
                if '{ContractAmountWithDiscountWords}' in run.text:
                    run.text = run.text.replace('{ContractAmountWithDiscountWords}',
                                                num2words(int(contract_amount_with_discount), lang="ru"))
                if '{ContractAmountWithDiscountWordsKaz}' in run.text:
                    run.text = run.text.replace('{ContractAmountWithDiscountWordsKaz}',
                                                num2words(int(contract_amount_with_discount), lang="kz"))
                if '{ContractAmountWithDiscountWordsEng}' in run.text:
                    run.text = run.text.replace('{ContractAmountWithDiscountWordsEng}',
                                                num2words(int(contract_amount_with_discount), lang="en"))
                if '{ContractDopAmountWords}' in run.text:
                    run.text = run.text.replace('{ContractDopAmountWords}',
                                                num2words(int(contract_dop_amount), lang="ru"))
                if '{ContractDopAmountWordsKaz}' in run.text:
                    run.text = run.text.replace('{ContractDopAmountWordsKaz}',
                                                num2words(int(contract_dop_amount), lang="kz"))
                if '{ContractContrWords}' in run.text:
                    contr_sum = int(contract.ContSum) if contract.ContSum else 0
                    if contr_sum > 0:
                        run.text = run.text.replace('{ContractContrWords}', num2words(contr_sum, lang="ru"))
                    else:
                        run.text = run.text.replace('{ContractContrWords}', 'ноль')
                if '{ContractContrWordsKaz}' in run.text:
                    contr_sum = int(contract.ContSum) if contract.ContSum else 0
                    if contr_sum > 0:
                        run.text = run.text.replace('{ContractContrWordsKaz}', num2words(contr_sum, lang="kz"))
                    else:
                        run.text = run.text.replace('{ContractContrWordsKaz}', 'нөл')
                if '{ContractContrWordsEng}' in run.text:
                    contr_sum = int(contract.ContSum) if contract.ContSum else 0
                    if contr_sum > 0:
                        run.text = run.text.replace('{ContractContrWordsEng}', num2words(contr_sum, lang="en"))
                    else:
                        run.text = run.text.replace('{ContractContrWordsEng}', 'zero')
            except:
                pass  # Игнорируем ошибки при преобразовании в слова

            # Тексты для QR-кодов и правовые тексты
            if '{QRCodeTextRus}' in run.text:
                run.text = run.text.replace('{QRCodeTextRus}',
                                            'QR-код содержит данные об электронно-цифровой подписи подписанта')
            if '{QRCodeTextKaz}' in run.text:
                run.text = run.text.replace('{QRCodeTextKaz}',
                                            'QR-кодта қол қоюшының электрондық-цифрлық қолтаңбасы туралы деректер қамтылады')
            if '{police_kaz}' in run.text:
                run.text = run.text.replace('{police_kaz}',
                                            'Осы құжат «Электрондық құжат және электрондық цифрлық қолтаңба туралы» Қазақстан Республикасының 2003 жылғы 7 қаңтардағы N 370-II Заңы 7 бабының 1 тармағына сәйкес қағаз тасығыштағы құжатпен бірдей.')
            if '{police_rus}' in run.text:
                run.text = run.text.replace('{police_rus}',
                                            'Данный документ согласно пункту 1 статьи 7 ЗРК от 7 января 2003 года «Об электронном документе и электронной цифровой подписи» равнозначен документу на бумажном носителе.')

        except Exception as e:
            logger.error(f"Error replacing text variables: {e}")

    def _replace_qr_codes(self, run, qr_signature, qr_director_omarov, qr_director_serikov, contract):
        """Заменяет QR-коды в run"""
        from io import BytesIO
        from docx.shared import Inches

        try:
            # Основные QR-коды подписи
            if '{QRCode}' in run.text or '{QRCodeSignature}' in run.text:
                run.text = run.text.replace('{QRCode}', '').replace('{QRCodeSignature}', '')
                if qr_signature:
                    image_stream = BytesIO(qr_signature)
                    run.add_picture(image_stream, width=Inches(1.5), height=Inches(1.5))

            # QR-код директора Омарова
            if '{QRCodeDirectorOmarov}' in run.text or '{QRcodeDirector}' in run.text:
                run.text = run.text.replace('{QRCodeDirectorOmarov}', '').replace('{QRcodeDirector}', '')
                if qr_director_omarov:
                    image_stream = BytesIO(qr_director_omarov)
                    run.add_picture(image_stream, width=Inches(1.5), height=Inches(1.5))

            # QR-код директора Серикова
            if '{QRCodeDirectorSerikov}' in run.text or '{QRCodeDirector2}' in run.text:
                run.text = run.text.replace('{QRCodeDirectorSerikov}', '').replace('{QRCodeDirector2}', '')
                if qr_director_serikov:
                    image_stream = BytesIO(qr_director_serikov)
                    run.add_picture(image_stream, width=Inches(1.5), height=Inches(1.5))

            # Специальные QR-коды
            if 'QRCodeDataSigned' in run.text:
                # Генерируем QR-код с подписанными данными
                qr_code_data_signed = self._generate_signed_data_qr_code(contract.ContractNum)
                if qr_code_data_signed:
                    qr_code_image = BytesIO(qr_code_data_signed)
                    run.add_picture(qr_code_image, width=Inches(1.3), height=Inches(1.3))
                    run.text = run.text.replace('QRCodeDataSigned', '')

        except Exception as e:
            logger.error(f"Error replacing QR codes: {e}")

    def _process_month_pay_table(self, cell, contract):
        """Обрабатывает таблицу помесячной оплаты"""
        try:
            from .models import ContractMonthPayMS, ContractDiscountMS, DiscountMS
            from docx.shared import Cm
            import math

            # Расчет суммы со скидкой
            discount = DiscountMS.objects.using('ms_sql').all()
            contract_discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract)
            contract_amount = float(contract.ContractAmount)

            if contract_discount.exists():
                for i in contract_discount:
                    for j in discount.filter(id=i.DiscountID.id):
                        percent = int(j.iDiscountPercent)
                        contract_amount -= contract_amount * percent / 100
                contract_amount_with_discount = float(round(contract_amount, 2))
            else:
                contract_amount_with_discount = float(round(contract_amount, 2))

            # Создаем таблицу
            main_table = cell.add_table(rows=1, cols=3)
            main_table.columns[0].width = Cm(1.0)
            main_table.columns[1].width = Cm(3.0)
            main_table.columns[2].width = Cm(3.0)

            # Заголовки
            header_row = main_table.rows[0]
            header_row.cells[0].text = '№'
            header_row.cells[1].text = 'Сумма'
            header_row.cells[2].text = 'Дата оплаты'

            # Получаем данные о платежах
            contract_month_pays = ContractMonthPayMS.objects.using('ms_sql').filter(ContractID=contract.id)
            count_month = 9
            sum_for_month = math.ceil(contract_amount_with_discount / count_month)

            # Добавляем строки
            for j, pay in enumerate(contract_month_pays):
                row_cells = main_table.add_row().cells
                row_cells[0].text = str(j + 1)
                row_cells[1].text = f"{sum_for_month:,}".replace(',', ' ')
                row_cells[2].text = str(pay.PayDateM) if pay.PayDateM else ''

        except Exception as e:
            logger.error(f"Error processing month pay table: {e}")

    def _process_quarter_pay_table(self, cell, contract):
        """Обрабатывает таблицу квартальной оплаты"""
        try:
            from .models import ContractMonthPayMS, ContractDiscountMS, DiscountMS
            from docx.shared import Cm
            from django.db.models import Sum, Min
            import math

            # Расчет суммы со скидкой
            discount = DiscountMS.objects.using('ms_sql').all()
            contract_discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract)
            contract_amount = float(contract.ContractAmount)

            if contract_discount.exists():
                for i in contract_discount:
                    for j in discount.filter(id=i.DiscountID.id):
                        percent = int(j.iDiscountPercent)
                        contract_amount -= contract_amount * percent / 100
                contract_amount_with_discount = float(round(contract_amount, 2))
            else:
                contract_amount_with_discount = float(round(contract_amount, 2))

            # Создаем таблицу
            main_table = cell.add_table(rows=1, cols=3)
            main_table.columns[0].width = Cm(1.0)
            main_table.columns[1].width = Cm(3.0)
            main_table.columns[2].width = Cm(3.0)

            # Заголовки
            header_row = main_table.rows[0]
            header_row.cells[0].text = '№'
            header_row.cells[1].text = 'Сумма'
            header_row.cells[2].text = 'Дата оплаты'

            # Получаем квартальные платежи
            contract_quarter_pays = ContractMonthPayMS.objects.using('ms_sql').filter(ContractID=contract.id)
            count_month = 9
            sum_for_month = math.ceil(contract_amount_with_discount / count_month)

            group_by_quarter = contract_quarter_pays.values('QuarterDig').annotate(
                MonthSum=Sum('MonthSum'), PayDateM=Min('PayDateM'))

            count_quarterdig = {}
            for quarter in group_by_quarter:
                count_quarterdig[quarter['QuarterDig']] = contract_quarter_pays.filter(
                    QuarterDig=quarter['QuarterDig']).count()

            # Добавляем строки
            for j, quarter in enumerate(group_by_quarter):
                quarter['MonthSum'] = round(sum_for_month * count_quarterdig[quarter['QuarterDig']], 2)
                row_cells = main_table.add_row().cells
                row_cells[0].text = str(j + 1)
                row_cells[1].text = str(quarter['MonthSum'])
                row_cells[2].text = str(quarter['PayDateM']) if quarter['PayDateM'] else ''

        except Exception as e:
            logger.error(f"Error processing quarter pay table: {e}")

    def _generate_signed_data_qr_code(self, contract_num):
        """Генерирует QR-код с данными подписанного договора"""
        try:
            from datetime import datetime

            qr_data = {
                "contract_num": contract_num,
                "signed_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "status": "Подписан",
                "verification_url": f"{self.frontend_url}/contracts/{contract_num}/signatures"
            }

            return self._create_qr_code(qr_data)

        except Exception as e:
            logger.error(f"Error generating signed data QR code: {e}")
            return b''

    def _get_contract_template(self, contract, is_dop_contract=False):
        """Получает шаблон контракта"""
        try:
            contract_payment_type = contract.PaymentTypeID.sPaymentType
            contract_school_language = getattr(contract.SchoolID, 'sSchool_language', '')
            contract_school_direct = getattr(contract.SchoolID, 'sSchool_direct', '')

            # Логика выбора шаблона из старого кода
            if contract_payment_type == 'Оплата по месячно':
                if contract_school_language == 'Казахское отделение':
                    if is_dop_contract:
                        return 'apps/contract/templates/contract/signed/Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_ТОО_по_месячно.docx'
                    else:
                        return 'apps/contract/templates/contract/signed/Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_по_месячно.docx'
                else:
                    if contract_school_direct == 'Кембридж':
                        return 'apps/contract/templates/contract/signed/Договор оказания образовательных услуг Кэмбридж 2025-2026 УО_по_месячно.docx'
                    elif is_dop_contract:
                        if contract_school_direct == 'Лингвинистический':
                            return 'apps/contract/templates/contract/signed/Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_по_месячно.docx'
                        elif contract_school_direct in ['Физико-математический', 'Физико-Математическая']:
                            return 'apps/contract/templates/contract/signed/Договор_оказания_дополнительных_образовательных_услуг_Физмат_Нур_по_месячно.docx'
                        elif contract_school_direct == 'Американская школа Advanced Placement':
                            return 'apps/contract/templates/contract/signed/Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_по_месячно.docx'
                        elif contract_school_direct == 'IT-школа на Кекилбайулы':
                            return 'apps/contract/templates/contract/signed/Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_по_месячно.docx'
                        else:
                            return 'apps/contract/templates/contract/signed/Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_месячно.docx'
                    else:
                        return 'apps/contract/templates/contract/signed/Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_месячно.docx'

            # Добавьте другие варианты оплаты по необходимости
            print("Default contract template used")
            logger.info("Default contract template used")
            return 'apps/contract/templates/contract/Договор оказания образовательных услуг_Школа 2025-2026_за_год.docx'

        except Exception as e:
            logger.error(f"Error getting contract template: {e}")
            return None

    def _generate_base_contract(self, contract, user, is_dop_contract=False):
        """Генерирует базовый контракт если файл не существует"""
        # Используем существующую логику из ChangeDocumentContentService
        from .services import ChangeDocumentContentService

        change_doc_service = ChangeDocumentContentService()
        student = contract.StudentID
        parent = student.parent_id

        change_doc_service.change_content(
            request=None,  # Можно передать None если не используется
            contract_num=contract.ContractNum,
            contract=contract,
            student=student,
            parent=parent,
            is_dop_contract=is_dop_contract
        )

    def _docx_to_pdf(self, input_path, output_path):
        """Конвертация файла из формата DOCX в PDF"""
        command_strings = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_path, input_path]
        try:
            subprocess.call(command_strings)
        except subprocess.CalledProcessError as e:
            logger.error(f"Error converting DOCX to PDF: {e}")

    def _calculate_contract_hash(self, contract, is_dop_contract=False) -> str:
        """Вычисляет хэш контракта на основе его ключевых данных"""
        contract_data = (
            f"{contract.ContractNum}:"
            f"{contract.ContractAmount}:"
            f"{contract.ContractDate}:"
            f"{getattr(contract, 'StudentID_id', '')}:"
            f"{getattr(contract, 'ContractStatusID_id', '')}"
        )

        try:
            if is_dop_contract:
                file_obj = ContractDopFileUser.objects.filter(contractNum=contract.ContractNum).first()
            else:
                file_obj = ContractFileUser.objects.filter(contractNum=contract.ContractNum).first()

            if file_obj and file_obj.file:
                hasher = hashlib.sha256()
                file_obj.file.seek(0)
                for chunk in file_obj.file.chunks():
                    hasher.update(chunk)
                file_hash = hasher.hexdigest()
                contract_data += f":{file_hash}"
        except Exception as e:
            logger.warning(f"Could not include file hash for contract {contract.ContractNum}: {e}")
            pass

        return hashlib.sha256(contract_data.encode()).hexdigest()

    def _add_director_signature(self, contract_num: str, parent_signature: ContractSignature, document_hash: str, signed_data):
        """Добавляет автоматическую подпись директора с правильным хэшем"""
        try:
            # Данные директора Омарова (заглушка - позже заполните реальными данными)
            director_omarov_data = {
                "iin": "540217301387",
                "full_name": "ОМАРОВ",
                "position": "Директор",
                "certificate_info": {
                    "valid": true,
                    "issuer": {
                        "dn": "C=KZ, CN=ҰЛТТЫҚ КУӘЛАНДЫРУШЫ ОРТАЛЫҚ (GOST) 2022",
                        "country": "KZ",
                        "commonName": "ҰЛТТЫҚ КУӘЛАНДЫРУШЫ ОРТАЛЫҚ (GOST) 2022"
                    },
                    "signAlg": "ECGOST3410-2015-512",
                    "subject": {
                        "dn": "GIVENNAME=СЫДЫКОВИЧ, OU=BIN070740004047, O=\"Товарищество с ограниченной ответственностью \\\"TAMOS EDUCATION\\\" (ТАМОС ЭДЬЮКЕЙШН)\", C=KZ, SERIALNUMBER=IIN540217301387, SURNAME=ОМАРОВ, CN=ОМАРОВ МУРАТ",
                        "bin": "070740004047",
                        "iin": "540217301387",
                        "country": "KZ",
                        "surName": "ОМАРОВ",
                        "commonName": "ОМАРОВ МУРАТ",
                        "organization": "Товарищество с ограниченной ответственностью \"TAMOS EDUCATION\" (ТАМОС ЭДЬЮКЕЙШН)"
                    },
                    "keyUsage": "SIGN",
                    "validity": {
                        "notAfter": "2026-05-19T04:50:54.000+00:00",
                        "notBefore": "2025-05-19T04:50:54.000+00:00"
                    },
                    "revocations": [
                        {
                            "by": "OCSP",
                            "reason": "OK",
                            "revoked": false,
                            "revocationTime": null
                        }
                    ],
                    "serialNumber": "59e4a35b6ac19e486926349d07eba7572059e0d8"
                }
            }
            director_serikov_data = {
                "iin": "861205300997",
                "full_name": "СЕРИКОВ",
                "position": "Директор",
                "certificate_info": {
                    "valid": true,
                    "issuer": {
                        "dn": "C=KZ, CN=ҰЛТТЫҚ КУӘЛАНДЫРУШЫ ОРТАЛЫҚ (GOST) 2022",
                        "country": "KZ",
                        "commonName": "ҰЛТТЫҚ КУӘЛАНДЫРУШЫ ОРТАЛЫҚ (GOST) 2022"
                    },
                    "signAlg": "ECGOST3410-2015-512",
                    "subject": {
                        "dn": "GIVENNAME=СЕРИКОВИЧ, OU=BIN990440006939, O=\"Учреждение образования \\\"Тамос Эдьюкейшн Физико-Математическая Школа\\\"\", C=KZ, SERIALNUMBER=IIN861205300997, SURNAME=СЕРИКОВ, CN=СЕРИКОВ БАУЫРЖАН",
                        "bin": "990440006939",
                        "iin": "861205300997",
                        "country": "KZ",
                        "surName": "СЕРИКОВ",
                        "commonName": "СЕРИКОВ БАУЫРЖАН",
                        "organization": "Учреждение образования \"Тамос Эдьюкейшн Физико-Математическая Школа\""
                    },
                    "keyUsage": "SIGN",
                    "validity": {
                        "notAfter": "2025-12-23T04:55:06.000+00:00",
                        "notBefore": "2024-12-23T04:55:06.000+00:00"
                    },
                    "revocations": [
                        {
                            "by": "OCSP",
                            "reason": "OK",
                            "revoked": false,
                            "revocationTime": null
                        }
                    ],
                    "serialNumber": "68492d0fd4cf74f7a417701044aef2663bff78a2"
                }
            }

            director_omarov_cms_signature = "-----BEGIN CMS-----\r\nMIIN2wYJKoZIhvcNAQcCoIINzDCCDcgCAQExDjAMBggqgw4DCgEDAwUAMAsGCSqG\r\nSIb3DQEHAaCCBPQwggTwMIIEWKADAgECAhRZ5KNbasGeSGkmNJ0H66dXIFng2DAO\r\nBgoqgw4DCgEBAgMCBQAwWDFJMEcGA1UEAwxA0rDQm9Ci0KLQq9KaINCa0KPTmNCb\r\n0JDQndCU0KvQoNCj0KjQqyDQntCg0KLQkNCb0KvSmiAoR09TVCkgMjAyMjELMAkG\r\nA1UEBhMCS1owHhcNMjUwNTE5MDQ1MDU0WhcNMjYwNTE5MDQ1MDU0WjCCASwxIDAe\r\nBgNVBAMMF9Ce0JzQkNCg0J7QkiDQnNCj0KDQkNCiMRUwEwYDVQQEDAzQntCc0JDQ\r\noNCe0JIxGDAWBgNVBAUTD0lJTjU0MDIxNzMwMTM4NzELMAkGA1UEBhMCS1oxgZIw\r\ngY8GA1UECgyBh9Ci0L7QstCw0YDQuNGJ0LXRgdGC0LLQviDRgSDQvtCz0YDQsNC9\r\n0LjRh9C10L3QvdC+0Lkg0L7RgtCy0LXRgtGB0YLQstC10L3QvdC+0YHRgtGM0Y4g\r\nIlRBTU9TIEVEVUNBVElPTiIgKNCi0JDQnNCe0KEg0K3QlNCs0K7QmtCV0JnQqNCd\r\nKTEYMBYGA1UECwwPQklOMDcwNzQwMDA0MDQ3MRswGQYDVQQqDBLQodCr0JTQq9Ca\r\n0J7QktCY0KcwgawwIwYJKoMOAwoBAQICMBYGCiqDDgMKAQECAgEGCCqDDgMKAQMD\r\nA4GEAASBgJlYKjvFkE6mo6bcImEtLNDzGBfNzqiWVoRIKjL7TY1aNJXMHJJQdZN5\r\nttkCLCL+z5qf2Prtg14W7exPtTVPc+bgX7AG7Y1tpONPkQas+BrJ1RNcKUWgdzPv\r\nMghLDV7w3PwLO29wi8aMaNxaPQ7/DOhoN2NV5dtkKMN2kV8QXIYfo4IB0DCCAcww\r\nDgYDVR0PAQH/BAQDAgPIMCgGA1UdJQQhMB8GCCsGAQUFBwMEBggqgw4DAwQBAgYJ\r\nKoMOAwMEAQIBMDgGA1UdIAQxMC8wLQYGKoMOAwMCMCMwIQYIKwYBBQUHAgEWFWh0\r\ndHA6Ly9wa2kuZ292Lmt6L2NwczA4BgNVHR8EMTAvMC2gK6AphidodHRwOi8vY3Js\r\nLnBraS5nb3Yua3ovbmNhX2dvc3RfMjAyMi5jcmwwOgYDVR0uBDMwMTAvoC2gK4Yp\r\naHR0cDovL2NybC5wa2kuZ292Lmt6L25jYV9kX2dvc3RfMjAyMi5jcmwwaAYIKwYB\r\nBQUHAQEEXDBaMCIGCCsGAQUFBzABhhZodHRwOi8vb2NzcC5wa2kuZ292Lmt6MDQG\r\nCCsGAQUFBzAChihodHRwOi8vcGtpLmdvdi5rei9jZXJ0L25jYV9nb3N0XzIwMjIu\r\nY2VyMB4GA1UdEQQXMBWBE3RhdHlhbmFfc2FkQG1haWwucnUwHQYDVR0OBBYEFNnk\r\no1tqwZ5IaSY0nQfrp1cgWeDYMB8GA1UdIwQYMBaAFP4wvp/IkGM/H/9aPAywyF9M\r\nbRcIMBYGBiqDDgMDBQQMMAoGCCqDDgMDBQEBMA4GCiqDDgMKAQECAwIFAAOBgQBi\r\nhf/b7tKvktNGOBoyFu2/P1c62CUMoe/2QrYg9rJc19gQCcKph94Zu/yvGohZ2VwT\r\nRkoogKSql5DgL2rwmOTrba8zloO3+aS6QMIemGU5NVSGk9+4LfYirRKIQ1FcQfbd\r\nADOXXT6K8PPOcGEbK0nZbeGWbR/Cs6Zw1BSNgGY3LTGCCKwwggioAgEBMHAwWDFJ\r\nMEcGA1UEAwxA0rDQm9Ci0KLQq9KaINCa0KPTmNCb0JDQndCU0KvQoNCj0KjQqyDQ\r\nntCg0KLQkNCb0KvSmiAoR09TVCkgMjAyMjELMAkGA1UEBhMCS1oCFFnko1tqwZ5I\r\naSY0nQfrp1cgWeDYMAwGCCqDDgMKAQMDBQCggcIwGAYJKoZIhvcNAQkDMQsGCSqG\r\nSIb3DQEHATAcBgkqhkiG9w0BCQUxDxcNMjUwOTE4MTgwNzI3WjA3BgsqhkiG9w0B\r\nCRACLzEoMCYwJDAiBCCl/pvrV/+nYoPfmz6t7B6ZFSxxv1vkvVz0UEeHHsHEdTBP\r\nBgkqhkiG9w0BCQQxQgRA5rn2Akys8YKZ1ms6fVHrG1wW1/JdtUkM3aOrrW1aeR/z\r\n3ok+XQ+qDqz3MC+cB6IZ6k32EImgAYt0fwzLsIBtJTAOBgoqgw4DCgEBAgMCBQAE\r\ngYBplIxsfNY7KMwNod4xjQzHGVIj83ZYANvgVCs81COMl+WHz83soCmgkTyFUzqp\r\nbdTZZdqAT9ekxD50PSUVN8JaH/7STifKs5GaBzZq/APaHa962PvMf19XysOdDrMW\r\nbgzp/Z1D3Q6ryEQddLWZu8bzfEi3Hh/vfifvBFnfnW7bLaGCBskwggbFBgsqhkiG\r\n9w0BCRACDjGCBrQwggawBgkqhkiG9w0BBwKgggahMIIGnQIBAzEOMAwGCCqDDgMK\r\nAQMDBQAwgaYGCyqGSIb3DQEJEAEEoIGWBIGTMIGQAgEBBggqgw4DAwIGBDBQMAwG\r\nCCqDDgMKAQMDBQAEQL7CQ6/alYAdIA6BJ1oSP1p5trtopInlm4raamuynM1a9h6j\r\nScPNiYnFc8jqUDts7oJS5DKDNd45gfZ3lFKC/pUCFE4CulX9yzQqfUgzWMCgghFW\r\nUsEHGA8yMDI1MDkxODE4MDcyN1oCCKbDCQI/evOmoIIEBDCCBAAwggNooAMCAQIC\r\nFBJ7KxdNTXWHNzZloR2fCDvbU6sjMA4GCiqDDgMKAQECAwIFADBYMUkwRwYDVQQD\r\nDEDSsNCb0KLQotCr0pog0JrQo9OY0JvQkNCd0JTQq9Cg0KPQqNCrINCe0KDQotCQ\r\n0JvQq9KaIChHT1NUKSAyMDIyMQswCQYDVQQGEwJLWjAeFw0yMjExMjYxOTAzMzVa\r\nFw0yNTExMjUxOTAzMzVaMG8xITAfBgNVBAMMGFRJTUUtU1RBTVBJTkcgQVVUSE9S\r\nSVRZCTELMAkGA1UEBhMCS1oxPTA7BgNVBAoMNNKw0JvQotCi0KvSmiDQmtCj05jQ\r\nm9CQ0J3QlNCr0KDQo9Co0Ksg0J7QoNCi0JDQm9Cr0powgawwIwYJKoMOAwoBAQIC\r\nMBYGCiqDDgMKAQECAgEGCCqDDgMKAQMDA4GEAASBgLKYaWKVHOsxLRpYzfvo091P\r\nSDR4azBDTAe7yzJFOUekA7WwfygIKWkBNEewRD20mfGZautmTx02O6yqngkc/5Bn\r\n2cnwmvSiK9sWzGwSmtyZLJ7p/9SYnsMLUJDM7yt0s0lQheH0fw61Vau0BB2bVj3r\r\n/MaYATnA+GmsOW2Rf7Yto4IBnzCCAZswFgYDVR0lAQH/BAwwCgYIKwYBBQUHAwgw\r\nOQYDVR0gBDIwMDAuBgcqgw4DAwIGMCMwIQYIKwYBBQUHAgEWFWh0dHA6Ly9wa2ku\r\nZ292Lmt6L2NwczBoBggrBgEFBQcBAQRcMFowIgYIKwYBBQUHMAGGFmh0dHA6Ly9v\r\nY3NwLnBraS5nb3Yua3owNAYIKwYBBQUHMAKGKGh0dHA6Ly9wa2kuZ292Lmt6L2Nl\r\ncnQvbmNhX2dvc3RfMjAyMi5jZXIwOAYDVR0fBDEwLzAtoCugKYYnaHR0cDovL2Ny\r\nbC5wa2kuZ292Lmt6L25jYV9nb3N0XzIwMjIuY3JsMDoGA1UdLgQzMDEwL6AtoCuG\r\nKWh0dHA6Ly9jcmwucGtpLmdvdi5rei9uY2FfZF9nb3N0XzIwMjIuY3JsMA4GA1Ud\r\nDwEB/wQEAwIHgDAdBgNVHQ4EFgQUknsrF01NdYc3NmWhHZ8IO9tTqyMwHwYDVR0j\r\nBBgwFoAU/jC+n8iQYz8f/1o8DLDIX0xtFwgwFgYGKoMOAwMFBAwwCgYIKoMOAwMF\r\nAQEwDgYKKoMOAwoBAQIDAgUAA4GBALdwN9n5WQda3OjIEieQu8BiSjMM55JdSJt0\r\nhSgay2YM1tXirYya5OcLcf8mD4xHZ5lLETbwxH4oPdMDePLpjudyvztsIa7YRpqC\r\n3p9ySSLn42kT2BXPP/zwYAbAn/QdZUc3nd4Ab0EE6jkSqN+g1jNDpl1TM0oNUBQw\r\nCe8eKyZ5MYIB1TCCAdECAQEwcDBYMUkwRwYDVQQDDEDSsNCb0KLQotCr0pog0JrQ\r\no9OY0JvQkNCd0JTQq9Cg0KPQqNCrINCe0KDQotCQ0JvQq9KaIChHT1NUKSAyMDIy\r\nMQswCQYDVQQGEwJLWgIUEnsrF01NdYc3NmWhHZ8IO9tTqyMwDAYIKoMOAwoBAwMF\r\nAKCBuDAaBgkqhkiG9w0BCQMxDQYLKoZIhvcNAQkQAQQwHAYJKoZIhvcNAQkFMQ8X\r\nDTI1MDkxODE4MDcyN1owKwYLKoZIhvcNAQkQAgwxHDAaMBgwFgQUlCxlK2qOUecZ\r\nygxFG8OVCyJwBIAwTwYJKoZIhvcNAQkEMUIEQMJsa6UOOu+QynThSZ/1jhZYb68Y\r\nMgEmPC8mYQMm6pLuRX/2uYU1GTRnlg/9pJCNXhFAVK45lzEsroz27OydvZcwDgYK\r\nKoMOAwoBAQIDAgUABIGAsJGTthvjPtIbU/Zd7KXrx6vu5Bj0s4UGcVg5f2EEdQgf\r\nre5qmRkarn4I58LKK46NJ9S/QkJ3gSiZRWIbw8sWN/JWyX3wuhubBC8Kuz6KO+Le\r\nlaT6c8BkZr6vHivnxhtBovBpHCYx9DTKVezKDX17kWec1Yy0IsvRuSU+096jbqs=\r\n-----END CMS-----\r\n",
            director_serikov_cms_signature = "-----BEGIN CMS-----\r\nMIIN7wYJKoZIhvcNAQcCoIIN4DCCDdwCAQExDjAMBggqgw4DCgEDAwUAMAsGCSqG\r\nSIb3DQEHAaCCBQgwggUEMIIEbKADAgECAhRoSS0P1M9096QXcBBErvJmO/94ojAO\r\nBgoqgw4DCgEBAgMCBQAwWDFJMEcGA1UEAwxA0rDQm9Ci0KLQq9KaINCa0KPTmNCb\r\n0JDQndCU0KvQoNCj0KjQqyDQntCg0KLQkNCb0KvSmiAoR09TVCkgMjAyMjELMAkG\r\nA1UEBhMCS1owHhcNMjQxMjIzMDQ1NTA2WhcNMjUxMjIzMDQ1NTA2WjCCAS8xKDAm\r\nBgNVBAMMH9Ch0JXQoNCY0JrQntCSINCR0JDQo9Cr0KDQltCQ0J0xFzAVBgNVBAQM\r\nDtCh0JXQoNCY0JrQntCSMRgwFgYDVQQFEw9JSU44NjEyMDUzMDA5OTcxCzAJBgNV\r\nBAYTAktaMYGLMIGIBgNVBAoMgYDQo9GH0YDQtdC20LTQtdC90LjQtSDQvtCx0YDQ\r\nsNC30L7QstCw0L3QuNGPICLQotCw0LzQvtGBINCt0LTRjNGO0LrQtdC50YjQvSDQ\r\npNC40LfQuNC60L4t0JzQsNGC0LXQvNCw0YLQuNGH0LXRgdC60LDRjyDQqNC60L7Q\r\nu9CwIjEYMBYGA1UECwwPQklOOTkwNDQwMDA2OTM5MRswGQYDVQQqDBLQodCV0KDQ\r\nmNCa0J7QktCY0KcwgawwIwYJKoMOAwoBAQICMBYGCiqDDgMKAQECAgEGCCqDDgMK\r\nAQMDA4GEAASBgLP3knoB6TTYN3MvOMnH2dCszqAWLLjuvYXTFRgXtUmycjdf7Ny9\r\nIsxEMkTEtvLiQabYSr8fnxNitmBj07vrPu1BHjXAhwUFbY924VTxKylPdkZiTrup\r\nqQjd7e4ekHy+4qhJPcdGBG1dm1qFBKrpX5pJ00Is5kun9tD55yuC1CYjo4IB4TCC\r\nAd0wDgYDVR0PAQH/BAQDAgPIMDIGA1UdJQQrMCkGCCqDDgMDBAMCBggrBgEFBQcD\r\nBAYIKoMOAwMEAQIGCSqDDgMDBAECATA4BgNVHSAEMTAvMC0GBiqDDgMDAjAjMCEG\r\nCCsGAQUFBwIBFhVodHRwOi8vcGtpLmdvdi5rei9jcHMwOAYDVR0fBDEwLzAtoCug\r\nKYYnaHR0cDovL2NybC5wa2kuZ292Lmt6L25jYV9nb3N0XzIwMjIuY3JsMGgGCCsG\r\nAQUFBwEBBFwwWjAiBggrBgEFBQcwAYYWaHR0cDovL29jc3AucGtpLmdvdi5rejA0\r\nBggrBgEFBQcwAoYoaHR0cDovL3BraS5nb3Yua3ovY2VydC9uY2FfZ29zdF8yMDIy\r\nLmNlcjA6BgNVHS4EMzAxMC+gLaArhilodHRwOi8vY3JsLnBraS5nb3Yua3ovbmNh\r\nX2RfZ29zdF8yMDIyLmNybDAlBgNVHREEHjAcgRpzZXJpa292YmF1eXJ6aGFuQGdt\r\nYWlsLmNvbTAdBgNVHQ4EFgQUaEktD9TPdPekF3AQRK7yZjv/eKIwHwYDVR0jBBgw\r\nFoAU/jC+n8iQYz8f/1o8DLDIX0xtFwgwFgYGKoMOAwMFBAwwCgYIKoMOAwMFAQEw\r\nDgYKKoMOAwoBAQIDAgUAA4GBAH46PTVn0ApnsXbUaLczMibB0Aeyu8uZxma4ofkQ\r\nmglgXEXom3ClnrMBf1PQBXDMcUwM0A8c1REcpKyBGUKSekYkD/BDcwz+ICDfXeV0\r\nA0Uy8Cy9qN44PYYKqlCQP9u9nYbskTk03wa8G83RoR4PVIyqSTiylstW64bC+9au\r\nlDh2MYIIrDCCCKgCAQEwcDBYMUkwRwYDVQQDDEDSsNCb0KLQotCr0pog0JrQo9OY\r\n0JvQkNCd0JTQq9Cg0KPQqNCrINCe0KDQotCQ0JvQq9KaIChHT1NUKSAyMDIyMQsw\r\nCQYDVQQGEwJLWgIUaEktD9TPdPekF3AQRK7yZjv/eKIwDAYIKoMOAwoBAwMFAKCB\r\nwjAYBgkqhkiG9w0BCQMxCwYJKoZIhvcNAQcBMBwGCSqGSIb3DQEJBTEPFw0yNTA5\r\nMTgxNzU4NThaMDcGCyqGSIb3DQEJEAIvMSgwJjAkMCIEINAxCVh3ClD6bKgn8xuU\r\nRNxL9ACAfCDqzfMQTIr4CUehME8GCSqGSIb3DQEJBDFCBEDmufYCTKzxgpnWazp9\r\nUesbXBbX8l21SQzdo6utbVp5H/PeiT5dD6oOrPcwL5wHohnqTfYQiaABi3R/DMuw\r\ngG0lMA4GCiqDDgMKAQECAwIFAASBgIPk/4BShhApK6AQ5iPRSX/JosZL1DdGzhel\r\nmt51eT3Nx87ylbvN06RKpTYDkjsDv3LXfZUReyV4mOaEDslYZHflArup/YA94vXb\r\nyGlbuyUcmNWcl+629KnZeR2APSVqhUqPDC7pXnyRuZaL2NRUi7T+xE3YJDhYCuRj\r\nq/tbuSbgoYIGyTCCBsUGCyqGSIb3DQEJEAIOMYIGtDCCBrAGCSqGSIb3DQEHAqCC\r\nBqEwggadAgEDMQ4wDAYIKoMOAwoBAwMFADCBpgYLKoZIhvcNAQkQAQSggZYEgZMw\r\ngZACAQEGCCqDDgMDAgYEMFAwDAYIKoMOAwoBAwMFAARAwiH6QOJAfq9/jPvGH+MW\r\n4+sSS1PRsCKZT9g4r8KnYMLIBr39HYWIcKGEyhYxeG9Hg8hNClmDwmA2zaA0o/ez\r\n8QIUypBNquOZnxrS0q84oLSLjbWobfMYDzIwMjUwOTE4MTc1ODU4WgIIuJhpawZu\r\nmn2gggQEMIIEADCCA2igAwIBAgIUEnsrF01NdYc3NmWhHZ8IO9tTqyMwDgYKKoMO\r\nAwoBAQIDAgUAMFgxSTBHBgNVBAMMQNKw0JvQotCi0KvSmiDQmtCj05jQm9CQ0J3Q\r\nlNCr0KDQo9Co0Ksg0J7QoNCi0JDQm9Cr0pogKEdPU1QpIDIwMjIxCzAJBgNVBAYT\r\nAktaMB4XDTIyMTEyNjE5MDMzNVoXDTI1MTEyNTE5MDMzNVowbzEhMB8GA1UEAwwY\r\nVElNRS1TVEFNUElORyBBVVRIT1JJVFkJMQswCQYDVQQGEwJLWjE9MDsGA1UECgw0\r\n0rDQm9Ci0KLQq9KaINCa0KPTmNCb0JDQndCU0KvQoNCj0KjQqyDQntCg0KLQkNCb\r\n0KvSmjCBrDAjBgkqgw4DCgEBAgIwFgYKKoMOAwoBAQICAQYIKoMOAwoBAwMDgYQA\r\nBIGAsphpYpUc6zEtGljN++jT3U9INHhrMENMB7vLMkU5R6QDtbB/KAgpaQE0R7BE\r\nPbSZ8Zlq62ZPHTY7rKqeCRz/kGfZyfCa9KIr2xbMbBKa3Jksnun/1JiewwtQkMzv\r\nK3SzSVCF4fR/DrVVq7QEHZtWPev8xpgBOcD4aaw5bZF/ti2jggGfMIIBmzAWBgNV\r\nHSUBAf8EDDAKBggrBgEFBQcDCDA5BgNVHSAEMjAwMC4GByqDDgMDAgYwIzAhBggr\r\nBgEFBQcCARYVaHR0cDovL3BraS5nb3Yua3ovY3BzMGgGCCsGAQUFBwEBBFwwWjAi\r\nBggrBgEFBQcwAYYWaHR0cDovL29jc3AucGtpLmdvdi5rejA0BggrBgEFBQcwAoYo\r\naHR0cDovL3BraS5nb3Yua3ovY2VydC9uY2FfZ29zdF8yMDIyLmNlcjA4BgNVHR8E\r\nMTAvMC2gK6AphidodHRwOi8vY3JsLnBraS5nb3Yua3ovbmNhX2dvc3RfMjAyMi5j\r\ncmwwOgYDVR0uBDMwMTAvoC2gK4YpaHR0cDovL2NybC5wa2kuZ292Lmt6L25jYV9k\r\nX2dvc3RfMjAyMi5jcmwwDgYDVR0PAQH/BAQDAgeAMB0GA1UdDgQWBBSSeysXTU11\r\nhzc2ZaEdnwg721OrIzAfBgNVHSMEGDAWgBT+ML6fyJBjPx//WjwMsMhfTG0XCDAW\r\nBgYqgw4DAwUEDDAKBggqgw4DAwUBATAOBgoqgw4DCgEBAgMCBQADgYEAt3A32flZ\r\nB1rc6MgSJ5C7wGJKMwznkl1Im3SFKBrLZgzW1eKtjJrk5wtx/yYPjEdnmUsRNvDE\r\nfig90wN48umO53K/O2whrthGmoLen3JJIufjaRPYFc8//PBgBsCf9B1lRzed3gBv\r\nQQTqORKo36DWM0OmXVMzSg1QFDAJ7x4rJnkxggHVMIIB0QIBATBwMFgxSTBHBgNV\r\nBAMMQNKw0JvQotCi0KvSmiDQmtCj05jQm9CQ0J3QlNCr0KDQo9Co0Ksg0J7QoNCi\r\n0JDQm9Cr0pogKEdPU1QpIDIwMjIxCzAJBgNVBAYTAktaAhQSeysXTU11hzc2ZaEd\r\nnwg721OrIzAMBggqgw4DCgEDAwUAoIG4MBoGCSqGSIb3DQEJAzENBgsqhkiG9w0B\r\nCRABBDAcBgkqhkiG9w0BCQUxDxcNMjUwOTE4MTc1ODU4WjArBgsqhkiG9w0BCRAC\r\nDDEcMBowGDAWBBSULGUrao5R5xnKDEUbw5ULInAEgDBPBgkqhkiG9w0BCQQxQgRA\r\nfyvnvijYw6EvVJL1aEolHVDjfgvyq7PZARtCcI1HiutvXFaEoVbv0jHGmyzdxV9k\r\nOtkHexxOs+JDB20Nje9u4zAOBgoqgw4DCgEBAgMCBQAEgYBxJcamT0aoFdH1uA5E\r\nXxC2y3hVjft+7gHBdtzqkH+lupOJqTm6fW6jo+csxcvls+loNVM35ijxR0JwrqV1\r\nP91Oz6tQx3zf9vC5jZoD5JC1w5GE23SEll4dukiKRua7HLvSAl2iOC8sBGOW8w25\r\nrcLFOunywr5XVpKxFILkO2irhA==\r\n-----END CMS-----\r\n"

            # Создаем подпись директора с ТЕМ ЖЕ хэшем что и у родителя
            director_signature = ContractSignature.objects.create(
                contract_num=contract_num,
                cms_signature=director_omarov_cms_signature,
                signed_data=signed_data,
                document_hash=document_hash,
                signer_iin=director_omarov_data["iin"],
                certificate_info=director_omarov_data["certificate_info"],
                is_valid=True,
                created_by=None
            )

            director_signature = ContractSignature.objects.create(
                contract_num=contract_num,
                cms_signature=director_serikov_cms_signature,
                signed_data=signed_data,
                document_hash=document_hash,
                signer_iin=director_serikov_data["iin"],
                certificate_info=director_serikov_data["certificate_info"],
                is_valid=True,
                created_by=None
            )

            logger.info(f"Director signature added for contract {contract_num} with hash {document_hash[:16]}...")

        except Exception as e:
            logger.error(f"Error adding director signature: {e}")

    # Остальные методы остаются такими же как в оригинальном коде...
    def get_contract_signatures(self, contract_num: str) -> Dict[str, Any]:
        """Получает все подписи для контракта"""
        try:
            # Проверяем существование контракта
            try:
                contract = ContractMS.objects.using('ms_sql').get(ContractNum=contract_num)
            except ContractMS.DoesNotExist:
                return {
                    'success': False,
                    'error': 'Контракт не найден',
                    'error_code': 'CONTRACT_NOT_FOUND'
                }

            # Получаем подписи по номеру контракта
            signatures = ContractSignature.objects.filter(
                contract_num=contract_num
            ).order_by('-signed_at')

            signatures_data = []
            for signature in signatures:
                signatures_data.append({
                    'signature_uid': str(signature.signature_uid),
                    'signer_iin': signature.signer_iin,
                    'signed_at': signature.signed_at.isoformat(),
                    'is_valid': signature.is_valid and not signature.is_document_modified,
                    'certificate_info': signature.certificate_info,
                    'is_document_modified': signature.is_document_modified
                })

            # Используем методы из ContractSignature
            signature_status = ContractSignature.get_signature_status(contract.ContractNum)

            return {
                'success': True,
                'contract_num': contract_num,
                'signature_status': signature_status,
                'signatures': signatures_data,
                'total_signatures': len(signatures_data),
                'valid_signatures': len([s for s in signatures_data if s['is_valid']])
            }

        except Exception as e:
            logger.error(f"Error in get_contract_signatures: {e}")
            return {
                'success': False,
                'error': f'Ошибка при получении подписей: {str(e)}',
                'error_code': 'PROCESSING_ERROR'
            }

    def check_signature_validity(self, signature_uid: str) -> Dict[str, Any]:
        """Проверяет актуальность подписи"""
        try:
            signature = ContractSignature.objects.get(signature_uid=signature_uid)

            # Проверяем не изменился ли документ
            is_document_modified = signature.is_document_modified

            if is_document_modified and signature.is_valid:
                # Помечаем подпись как невалидную если документ изменился
                signature.is_valid = False
                signature.save()
                logger.info(f"Signature {signature_uid} marked as invalid due to document modification")

            return {
                'success': True,
                'signature_uid': signature_uid,
                'contract_num': signature.contract_num,
                'is_valid': signature.is_valid,
                'is_document_modified': is_document_modified,
                'signed_at': signature.signed_at.isoformat(),
                'signer_iin': signature.signer_iin
            }

        except ContractSignature.DoesNotExist:
            return {
                'success': False,
                'error': 'Подпись не найдена',
                'error_code': 'SIGNATURE_NOT_FOUND'
            }
        except Exception as e:
            logger.error(f"Error in check_signature_validity: {e}")
            return {
                'success': False,
                'error': f'Ошибка при проверке подписи: {str(e)}',
                'error_code': 'PROCESSING_ERROR'
            }

    def get_contract_summary(self, contract_num: str) -> Dict[str, Any]:
        """Получает краткую сводку по контракту и его подписям"""
        try:
            contract = ContractMS.objects.using('ms_sql').get(ContractNum=contract_num)
            signatures = ContractSignature.get_contract_signatures(contract_num)

            valid_signatures = signatures.filter(is_valid=True)

            return {
                'success': True,
                'contract_num': contract_num,
                'contract_info': {
                    'student_name': getattr(contract.StudentID, 'full_name', '') if hasattr(contract,
                                                                                            'StudentID') and contract.StudentID else '',
                    'contract_amount': str(contract.ContractAmount) if contract.ContractAmount else '',
                    'contract_date': contract.ContractDate.isoformat() if contract.ContractDate else '',
                },
                'signature_summary': {
                    'total_signatures': signatures.count(),
                    'valid_signatures': valid_signatures.count(),
                    'status': ContractSignature.get_signature_status(contract.ContractNum),
                    'last_signed': signatures.first().signed_at.isoformat() if signatures.exists() else None,
                    'has_valid_signatures': valid_signatures.exists()
                }
            }

        except ContractMS.DoesNotExist:
            return {
                'success': False,
                'error': 'Контракт не найден',
                'error_code': 'CONTRACT_NOT_FOUND'
            }
        except Exception as e:
            logger.error(f"Error in get_contract_summary: {e}")
            return {
                'success': False,
                'error': f'Ошибка при получении сводки: {str(e)}',
                'error_code': 'PROCESSING_ERROR'
            }