import math
import os
import subprocess
from decimal import Decimal
from datetime import datetime, timedelta
from io import BytesIO

from django.core.exceptions import ObjectDoesNotExist, MultipleObjectsReturned
from django.core.files.base import ContentFile
from django.db.models import Sum, Min
from django.http import FileResponse, JsonResponse
from docx.shared import Inches, Cm
from translate import Translator
from docx import Document

from num2words import num2words

from rest_framework import status
from rest_framework.response import Response

from .models import TransactionMS, ContractMonthPayMS, ContractFoodMonthPayMS, TransactionFoodMS, \
    ContractDriverMonthPayMS, TransactionDriverMS, ParentMS, StudentMS, ContractFileUser, ContractDiscountMS, \
    ContractFoodDiscountMS, ContractDopFileUser, ContractDopMS, DiscountMS, ContractMS
from .serializers.contract import ContractSerializer
from .serializers.contract_driver import ContractDriverSerializer
from .serializers.contract_food import ContractFoodSerializer


class GetQuerySet:
    """ Фильтрация данных по переданному аргументу """

    @staticmethod
    def get_queryset(model, filter_argument, filter_value):
        return model.objects.using('ms_sql').filter(**{filter_argument: filter_value})


class ContractService:
    """
        Сервис для получения данных о договоре студента.
        Договоров может быть несколько, поэтому возвращается список.
        Договор ищется по ID студента.
    """

    def __init__(self, contract) -> None:
        self.contract = contract

    @staticmethod
    def get_queryset(model, filter_argument, filter_value):
        """ Фильтрация данных по переданному аргументу """
        return model.objects.using('ms_sql').filter(**{filter_argument: filter_value})

    def sum_month_pay(self, contract_id=None) -> float:
        sum_month_pay = self.get_queryset(ContractMonthPayMS, 'ContractID', contract_id)
        amount = 0
        for sum_month in sum_month_pay:
            amount += sum_month.MonthSum

        return float(amount)

    def sum_transactions_without_contribution(self, contract_id=None) -> float:
        sum_transactions = self.get_queryset(TransactionMS, 'agreement_id', contract_id)
        amount = 0
        for sum_transaction in sum_transactions:
            if sum_transaction.contribution == 0:
                amount += sum_transaction.amount

        return float(amount)

    def sum_transactions_month_pay(self, contract_student_filter=None) -> list:
        sum_transactions_list = list()
        for contract_student in contract_student_filter:
            agreement_id = contract_student.id
            if agreement_id in sum_transactions_list:
                continue

            transaction = self.get_queryset(TransactionMS, 'agreement_id', agreement_id)
            if transaction.exists():
                sum_month_pay = self.sum_month_pay(contract_id=agreement_id)
                sum_transactions_list.append((f'{agreement_id}', sum_month_pay))

        return sum_transactions_list

    def calculate_arrears(self, contract_num):
        # Получите контракт по номеру
        contract = self.contract.filter(ContractNum=contract_num).first()
        if not contract:
            return Decimal(0)  # Возвращайте 0, если контракт не найден

            # Получите общую сумму платежей за контракт и преобразуйте ее в Decimal
        total_payments = Decimal(self.sum_month_pay(contract.id))

        # Получите общую сумму транзакций без взносов за контракт и преобразуйте ее в Decimal
        total_transactions = Decimal(self.sum_transactions_without_contribution(contract.id))

        # Вычислите задолженность как разницу между суммой контракта и общей суммой платежей и транзакций
        arrears = contract.ContractSum - (total_payments + total_transactions)

        return arrears

    def get_value_of_arrears_with_contract_num(self, contract_num):
        contract = self.contract
        contract_student_filter = contract.filter(ContractNum=contract_num)
        serializer = ContractSerializer(contract_student_filter, many=True)

        # Вычислите задолженность и добавьте ее к данным договора
        arrears = self.calculate_arrears(contract_num)  # Замените на вашу функцию для вычисления задолженности
        serializer_data = serializer.data
        if serializer_data:
            serializer_data[0]['Arrears'] = arrears

        return serializer_data[0]

    def set_arrears_from_sum_transactions(self, contract_student_filter=None, serializer_data=None) -> None:
        sum_transactions_list = self.sum_transactions_month_pay(contract_student_filter=contract_student_filter)

        for ser in serializer_data:
            for transactions in sum_transactions_list:
                sum_transactions = float(transactions[1])

                if ser['id'] == int(transactions[0]):
                    if sum_transactions > 0:
                        arrears_value = round(float(sum_transactions) - self.sum_transactions_without_contribution(contract_id=transactions[0]), 2)

                        if arrears_value < 1:
                            arrears_value = 0
                        ser['Arrears'] = math.ceil(arrears_value)
                    else:
                        ser['Arrears'] = 0

    def set_discount_to_contract(self, contract_student_filter=None, serializer_data=None) -> None:
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        discount = ContractDiscountMS.objects.using('ms_sql').get(ContractID=contract_student.id)

                        discount_data = {
                            "DiscountID": str(discount.DiscountID),
                            "DiscountName": str(discount.DiscountID.sDiscountName),
                            "DiscountPercent": str(discount.DiscountID.iDiscountPercent),
                            "DiscountType": str(discount.DiscountID.iDiscountType.sDiscountType),
                            "DiscountSum": str(discount.DiscountSum)
                        }

                        ser['Discount'] = dict(discount_data)
                    except ObjectDoesNotExist:
                        ser['Discount'] = None
                    except MultipleObjectsReturned:
                        discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract_student.id).last()

                        discount_data = {
                            "DiscountID": str(discount.DiscountID),
                            "DiscountName": str(discount.DiscountID.sDiscountName),
                            "DiscountPercent": str(discount.DiscountID.iDiscountPercent),
                            "DiscountType": str(discount.DiscountID.iDiscountType.sDiscountType),
                            "DiscountSum": str(discount.DiscountSum)
                        }

                        ser['Discount'] = dict(discount_data)

    @staticmethod
    def set_history_of_month_pays(contract_student_filter=None, serializer_data=None):
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        history_pays = TransactionMS.objects.using('ms_sql').filter(agreement_id=contract_student.id)
                        transactions = []

                        for pay in history_pays:
                            transactions.append({
                                "Amount": str(pay.amount),
                                "Description": str(pay.description),
                                "Date": str(pay.trans_date.strftime("%d.%m.%Y")),
                                "Bank": str(pay.bank_id),
                                "PaymentType": str(pay.payment_type.sPaymentType)
                            })

                        ser['HistoryTransactions'] = list(transactions)
                    except ObjectDoesNotExist:
                        ser['HistoryTransactions'] = []

    def get_value_of_arrears(self, contract_num):
        try:
            contract = self.contract
            contract_student_filter = contract.filter(ContractNum=contract_num)
            serializer = ContractSerializer(contract_student_filter, many=True)
            self.set_arrears_from_sum_transactions(
                contract_student_filter=contract_student_filter,
                serializer_data=serializer.data
            )
            return serializer.data[0]['Arrears']
        except IndexError:
            return None
        except KeyError:
            return None

    def get_contract(self, student_id=None) -> Response | JsonResponse:
        contract = self.contract
        contract_student_filter = contract.filter(StudentID=student_id)
        print("CONTRACT:", contract_student_filter)

        if not contract_student_filter.exists():
            return Response({'error': 'Contract not found'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ContractSerializer(contract_student_filter, many=True)
        self.set_arrears_from_sum_transactions(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        self.set_discount_to_contract(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        self.set_history_of_month_pays(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )

        return JsonResponse(serializer.data, safe=False)


class ContractFoodService:
    def __init__(self, contract) -> None:
        self.contract = contract

    @staticmethod
    def get_queryset(model, filter_argument, filter_value):
        """ Фильтрация данных по переданному аргументу """
        return model.objects.using('ms_sql').filter(**{filter_argument: filter_value})

    def sum_month_pay(self, contract_id=None) -> float:
        sum_month_pay = self.get_queryset(ContractFoodMonthPayMS, 'ContractID', contract_id)
        amount = 0
        for sum_month in sum_month_pay:
            amount += sum_month.MonthSum

        return float(amount)

    def sum_transactions_without_contribution(self, contract_id=None) -> float:
        sum_transactions = self.get_queryset(TransactionFoodMS, 'contract_id', contract_id)
        amount = 0
        for sum_transaction in sum_transactions:
            amount += sum_transaction.amount

        return float(amount)

    def sum_transactions_month_pay(self, contract_student_filter=None) -> list:
        sum_transactions_list = list()
        for contract_student in contract_student_filter:
            agreement_id = contract_student.id
            if agreement_id in sum_transactions_list:
                continue

            transaction = self.get_queryset(TransactionFoodMS, 'contract_id', agreement_id)
            if transaction.exists():
                sum_month_pay = self.sum_month_pay(contract_id=agreement_id)
                sum_transactions_list.append((f'{agreement_id}', sum_month_pay))

        return sum_transactions_list

    def get_value_of_arrears(self, contract_num):
        contract = self.contract
        contract_student_filter = contract.filter(ContractNum=contract_num)
        serializer = ContractSerializer(contract_student_filter, many=True)
        self.set_arrears_from_sum_transactions(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        return serializer.data[0]['Arrears']

    def set_arrears_from_sum_transactions(self, contract_student_filter=None, serializer_data=None) -> None:
        sum_transactions_list = self.sum_transactions_month_pay(contract_student_filter=contract_student_filter)

        for ser in serializer_data:
            for transactions in sum_transactions_list:
                sum_transactions = float(transactions[1])

                if ser['id'] == int(transactions[0]):
                    if sum_transactions > 0:
                        arrears_value = round(float(sum_transactions) - self.sum_transactions_without_contribution(
                            contract_id=transactions[0]), 2)

                        if arrears_value < 1:
                            arrears_value = 0
                        ser['Arrears'] = math.ceil(arrears_value)
                    else:
                        ser['Arrears'] = 0

    def set_discount_to_contract_food(self, contract_student_filter=None, serializer_data=None) -> None:
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        discount = ContractFoodDiscountMS.objects.using('ms_sql').get(ContractID=contract_student.id)

                        discount_data = {
                            "DiscountID": str(discount.DiscountID),
                            "DiscountName": str(discount.DiscountID.sDiscountName),
                            "DiscountPercent": str(discount.DiscountID.iDiscountPercent),
                            "DiscountType": str(discount.DiscountID.iDiscountType.sDiscountType),
                            "DiscountSum": str(discount.DiscountSum)
                        }

                        ser['Discount'] = dict(discount_data)
                    except ObjectDoesNotExist:
                        ser['Discount'] = None

    @staticmethod
    def set_history_of_month_pays_food(contract_student_filter=None, serializer_data=None):
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        history_pays = TransactionFoodMS.objects.using('ms_sql').filter(contract_id=contract_student.id)
                        transactions = []

                        for pay in history_pays:
                            transactions.append({
                                "Amount": str(pay.amount),
                                "Description": str(pay.description),
                                "Date": str(pay.trans_date.strftime("%d.%m.%Y")),
                                "Bank": str(pay.bank_id)
                            })

                        ser['HistoryTransactions'] = list(transactions)
                    except ObjectDoesNotExist:
                        ser['HistoryTransactions'] = []

    @staticmethod
    def set_detail_contract_food(contract_student_filter=None, serializer_data=None):
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        pays = ContractFoodMonthPayMS.objects.using('ms_sql').filter(ContractID=contract_student.id)
                        month_pays = []
                        for pay in pays:
                            month_pays.append({
                                "MonthAmount": str(pay.MonthAmount),
                                "MonthSum": str(pay.MonthSum),
                                "PayDateM": str(pay.PayDateM.strftime("%d.%m.%Y"))
                            })

                        ser['DetailContract'] = list(month_pays)
                    except ObjectDoesNotExist:
                        ser['DetailContract'] = []

    def get_contract_food(self, student_id=None) -> Response:
        contract = self.contract
        contract_student_filter = contract.filter(StudentID=student_id)

        if not contract_student_filter.exists():
            return Response({'error': 'Contract not found'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ContractFoodSerializer(contract_student_filter, many=True)
        self.set_arrears_from_sum_transactions(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        self.set_discount_to_contract_food(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        self.set_history_of_month_pays_food(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )
        self.set_detail_contract_food(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )

        return Response(serializer.data)


class ContractDriverService:
    def __init__(self, contract) -> None:
        self.contract = contract

    @staticmethod
    def get_queryset(model, filter_argument, filter_value):
        """ Фильтрация данных по переданному аргументу """
        return model.objects.using('ms_sql').filter(**{filter_argument: filter_value})

    def sum_month_pay(self, contract_id=None) -> float:
        sum_month_pay = self.get_queryset(ContractDriverMonthPayMS, 'ContractID', contract_id)
        amount = 0
        for sum_month in sum_month_pay:
            amount += sum_month.MonthAmount

        return float(amount)

    def sum_transactions_without_contribution(self, contract_id=None) -> float:
        sum_transactions = self.get_queryset(TransactionDriverMS, 'ContractID', contract_id)
        amount = 0
        for sum_transaction in sum_transactions:
            amount += sum_transaction.Amount

        return float(amount)

    def sum_transactions_month_pay(self, contract_student_filter=None) -> list:
        sum_transactions_list = list()
        for contract_student in contract_student_filter:
            agreement_id = contract_student.id
            if agreement_id in sum_transactions_list:
                continue

            transaction = self.get_queryset(TransactionDriverMS, 'ContractID', agreement_id)
            if transaction.exists():
                sum_month_pay = self.sum_month_pay(contract_id=agreement_id)
                sum_transactions_list.append((f'{agreement_id}', sum_month_pay))

        return sum_transactions_list

    def get_value_of_arrears(self, contract_num):
        contract = self.contract
        contract_student_filter = contract.filter(ContractNum=contract_num)
        serializer = ContractSerializer(contract_student_filter, many=True)
        self.set_arrears_from_sum_transactions(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )

        try:
            return serializer.data[0]['Arrears']
        except KeyError:
            return 0

    def set_arrears_from_sum_transactions(self, contract_student_filter=None, serializer_data=None) -> None:
        sum_transactions_list = self.sum_transactions_month_pay(contract_student_filter=contract_student_filter)

        for ser in serializer_data:
            for transactions in sum_transactions_list:
                sum_transactions = float(transactions[1])

                if ser['id'] == int(transactions[0]):
                    if sum_transactions > 0:
                        arrears_value = round(float(sum_transactions) - self.sum_transactions_without_contribution(
                            contract_id=transactions[0]), 2)

                        if arrears_value < 1:
                            arrears_value = 0
                        ser['Arrears'] = math.ceil(arrears_value)
                    else:
                        ser['Arrears'] = 0

    @staticmethod
    def set_history_of_month_pays_driver(contract_student_filter=None, serializer_data=None):
        for ser in serializer_data:
            for contract_student in contract_student_filter:
                if ser['id'] == contract_student.id:
                    try:
                        history_pays = TransactionDriverMS.objects.using('ms_sql').filter(ContractID=contract_student.id)
                        transactions = []

                        for pay in history_pays:
                            transactions.append({
                                "Amount": str(pay.Amount),
                                "Description": str(pay.Description),
                                "Date": str(pay.TransactionDate.strftime("%d.%m.%Y")),
                                "Bank": str(pay.BankID)
                            })

                        ser['HistoryTransactions'] = list(transactions)
                    except ObjectDoesNotExist:
                        ser['HistoryTransactions'] = []

    def get_contract_driver(self, student_id=None) -> Response:
        contract = self.contract
        contract_student_filter = contract.filter(StudentID=student_id)

        if not contract_student_filter.exists():
            return Response({'error': 'Contract not found'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ContractDriverSerializer(contract_student_filter, many=True)
        self.set_arrears_from_sum_transactions(
            contract_student_filter=contract_student_filter,
            serializer_data=serializer.data
        )

        return Response(serializer.data)


class GetContractFromDBService:
    """ Получение договора по номеру из базы данных """

    @staticmethod
    def get_contract(contract_num, is_dop_contract):
        if is_dop_contract:
            try:
                contract = ContractDopFileUser.objects.get(contractNum=contract_num)
            except ContractDopFileUser.DoesNotExist:
                contract = None
            except ContractDopFileUser.MultipleObjectsReturned:
                contract = ContractDopFileUser.objects.filter(contractNum=contract_num).last()
        else:
            try:
                contract = ContractFileUser.objects.get(contractNum=contract_num)
            except ContractFileUser.DoesNotExist:
                contract = None
            except ContractFileUser.MultipleObjectsReturned:
                contract = ContractFileUser.objects.filter(contractNum=contract_num).last()

        if contract is not None:
            return contract

        return None


class ChangeDocumentContentService:
    """
        Сервис для изменения содержимого документа.
        Документ ищется по номеру договора.
        Содержимое документа изменяется в соответствии с переданными данными.
    """

    @staticmethod
    def translate_text(text, dest_lang):
        translator = Translator(to_lang=dest_lang)
        translation = translator.translate(text)
        return translation

    @staticmethod
    def docx_to_pdf(input_path, output_path):
        """ Конвертация файла из формата DOCX в PDF """

        command_strings = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_path, input_path]

        try:
            subprocess.call(command_strings)
        except subprocess.CalledProcessError as e:
            return Response({'error': str(e)}, status=status.HTTP_403_FORBIDDEN)

    def save_document(self, request, contract_num, doc, docx_output_path, pdf_directory, pdf_output_path, is_dop_contract):
        doc.save(docx_output_path)
        self.docx_to_pdf(docx_output_path, pdf_directory)

        # os.remove(docx_output_path)

        try:
            with open(pdf_output_path, 'rb') as pdf_file:
                file_content = pdf_file.read()
        except FileNotFoundError:
            print('Файл не найден!')
            return {'error': 'Для подписания договора файл не найден! Возможно, указали не верный номер контракта'}

        except Exception as e:
            print(f'Ошибка при чтении файла: {e}')
            return {'error': str(e)}

        if is_dop_contract:
            try:
                contract_file_user = ContractDopFileUser.objects.create(
                    user=request.user,
                    contractNum=contract_num,
                    file=ContentFile(file_content, name=f'{contract_num}.pdf')
                )
            except Exception as e:
                print(f'Exception Dop Contract: {e}')
                contract_file_user = ContractDopFileUser.objects.create(
                    user=None,
                    contractNum=contract_num,
                    file=ContentFile(file_content, name=f'{contract_num}.pdf')
                )
        else:
            try:
                contract_file_user = ContractFileUser.objects.create(
                    user=request.user,
                    contractNum=contract_num,
                    file=ContentFile(file_content, name=f'{contract_num}.pdf')
                )
            except Exception as e:
                print(f'Exception Contract: {e}')
                contract_file_user = ContractFileUser.objects.create(
                    user=None,
                    contractNum=contract_num,
                    file=ContentFile(file_content, name=f'{contract_num}.pdf')
                )

        return contract_file_user.file

    @staticmethod
    def get_contract_type_month_pay(contract_id):
        contract_month_pay = ContractMonthPayMS.objects.using('ms_sql').filter(ContractID=contract_id)

        if contract_month_pay.exists():
            return contract_month_pay
        else:
            return None

    def change_docx_contract_month_pay(self, doc, contract, is_dop_contract):
        if is_dop_contract:
            try:
                dop_contract = ContractMS.objects.using('ms_sql').filter(ContractNum=contract.ContractNum).first()
                contract_amount_with_discount = dop_contract.ContractAmount
                if '-' in str(contract_amount_with_discount):
                    contract_amount_with_discount = str(contract_amount_with_discount).replace('-', '')
                contract_amount_with_discount = int(float(contract_amount_with_discount))
            except ObjectDoesNotExist:
                contract_amount_with_discount = 0
        else:
            discount = DiscountMS.objects.using('ms_sql').all()
            contract_discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract)
            contract_amount = float(contract.ContractAmount)

            if contract_discount.exists():
                for i in contract_discount:
                    for j in discount.filter(id=i.DiscountID.id):
                        percent = int(j.iDiscountPercent)
                        contract_amount -= contract_amount * percent / 100

                contract_amount_with_discount = float(round(contract_amount, 2))
            else:
                contract_amount_with_discount = float(round(contract_amount, 2))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        inline = paragraph
                        for i in inline.runs:
                            count_month = 9
                            sum_for_month = math.ceil(int(contract_amount_with_discount) / count_month)

                            if '{customtable_monthpay}' in i.text:
                                main_table = cell.add_table(rows=1, cols=3)

                                main_table.columns[0].width = Cm(1.0)
                                main_table.columns[1].width = Cm(3.0)
                                main_table.columns[2].width = Cm(3.0)

                                # main_table.style = 'Table Grid'

                                header_row = main_table.rows[0]
                                header_row.cells[0].text = '№'
                                header_row.cells[1].text = 'Сумма'
                                header_row.cells[2].text = 'Дата оплаты'

                                try:
                                    contract_month_pays = self.get_contract_type_month_pay(contract.id)
                                except ObjectDoesNotExist:
                                    contract_month_pays = None
                                    sum_for_month = 0

                                for j in range(len(contract_month_pays)):
                                    row_cells = main_table.add_row().cells
                                    row_cells[0].text = str(j + 1)
                                    row_cells[1].text = f"{sum_for_month:,}".replace(',', ' ')
                                    row_cells[2].text = str(contract_month_pays[j].PayDateM)

                                text = i.text.replace('{customtable_monthpay}', '')
                                i.text = text

                                paragraph._p.addnext(main_table._tbl)
                                paragraph._p.getparent().remove(paragraph._p)

                            if '{customtable_quarterpay}' in i.text:
                                main_table = cell.add_table(rows=1, cols=3)

                                main_table.columns[0].width = Cm(1.0)
                                main_table.columns[1].width = Cm(3.0)
                                main_table.columns[2].width = Cm(3.0)

                                # main_table.style = 'Table Grid'

                                header_row = main_table.rows[0]
                                header_row.cells[0].text = '№'
                                header_row.cells[1].text = 'Сумма'
                                header_row.cells[2].text = 'Дата оплаты'

                                try:
                                    contract_quarter_pays = self.get_contract_type_month_pay(contract.id)
                                except AttributeError:
                                    raise AttributeError('Contract has no quarter pays')
                                try:
                                    group_by_quarter = contract_quarter_pays.values('QuarterDig').annotate(
                                        MonthSum=Sum('MonthSum'), PayDateM=Min('PayDateM'))
                                    count_quarterdig = {}

                                    for j in range(len(group_by_quarter)):
                                        """ Количество месяцев одинаковых кварталов """

                                        count_quarterdig[group_by_quarter[j]['QuarterDig']] = \
                                            contract_quarter_pays.filter(
                                                QuarterDig=group_by_quarter[j]['QuarterDig']).count()

                                except AttributeError:
                                    raise AttributeError('Contract has no quarter pays')

                                for j in range(len(group_by_quarter)):
                                    group_by_quarter[j]['MonthSum'] = round(
                                        sum_for_month * count_quarterdig[group_by_quarter[j]['QuarterDig']], 2)
                                    row_cells = main_table.add_row().cells
                                    row_cells[0].text = str(j + 1)
                                    row_cells[1].text = str(group_by_quarter[j]['MonthSum'])
                                    row_cells[2].text = str(group_by_quarter[j]['PayDateM'])

                                text = i.text.replace('{customtable_quarterpay}', '')
                                i.text = text

                                paragraph._p.addnext(main_table._tbl)
                                paragraph._p.getparent().remove(paragraph._p)

    def change_docx_document(self, doc, contract, student, parent, is_dop_contract):
        ParentPassport = f'Удостоверение личности: №{parent.num_of_doc}, Орган выдачи: {parent.issued_by}, Дата выдачи: {parent.issue_date}'
        ParentPassportKAZ = f'Жеке куәлік: №{parent.num_of_doc}, Берген орган: {parent.issued_by}, Берілген күні: {parent.issue_date}'
        ParentPassportENG = f'ID: No.{parent.num_of_doc}, Issued by: {parent.issued_by}, Issue date: {parent.issue_date}'
        if '?' in str(parent.issued_by):
            ParentPassport = f'Удостоверение личности: №{parent.num_of_doc}, Орган выдачи: Не указано, Дата выдачи: {parent.issue_date}'
            ParentPassportKAZ = f'Жеке куәлік: №{parent.num_of_doc}, Берген орган: Көрсетілмеген, Берілген күні: {parent.issue_date}'
            ParentPassportENG = f'ID: No.{parent.num_of_doc}, Issued by: Not specified, Issue date: {parent.issue_date}'

        contract_dop_amount = ContractMS.objects.using('ms_sql').filter(ContractNum=contract.ContractNum).first()
        if contract_dop_amount is not None:
            contract_dop_amount = contract_dop_amount.ContractAmount
        else:
            contract_dop_amount = 0

        discount = DiscountMS.objects.using('ms_sql').all()
        contract_discount = ContractDiscountMS.objects.using('ms_sql').filter(ContractID=contract)

        contract_amount = float(contract.ContractAmount)

        if contract_discount.exists():
            for i in contract_discount:
                for j in discount.filter(id=i.DiscountID.id):
                    percent = int(j.iDiscountPercent)
                    contract_amount -= contract_amount * percent / 100

            contract_amount_with_discount = float(round(contract_amount, 2))
        else:
            contract_amount_with_discount = float(round(contract_amount, 2))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        inline = paragraph
                        for i in inline.runs:
                            if '{ContractNum}' in i.text:
                                text = i.text.replace('{ContractNum}', f'{contract.ContractNum}')
                                i.text = text
                            if '{ContractYear}' in i.text:
                                text = i.text.replace('{ContractYear}', f'{contract.ContractDate.strftime("%Y")}')
                                i.text = text
                            if '{ContractYearFinish}' in i.text:
                                data_close = contract.ContractDate + timedelta(days=365)
                                text = i.text.replace('{ContractYearFinish}', f'{data_close.strftime("%Y")}')
                                i.text = text
                            if '{ContractDate}' in i.text:
                                text = i.text.replace('{ContractDate}', f'{contract.ContractDate.strftime("%d.%m.%Y")}')
                                i.text = text
                            if '{ContractDay}' in i.text:
                                text = i.text.replace('{ContractDay}', f'{contract.ContractDate.strftime("%d")}')
                                i.text = text
                            if '{ContractMonthRUS}' in i.text:
                                text = i.text.replace('{ContractMonthRUS}', f'{self.translate_text(contract.ContractDate.strftime("%B"), "ru")}')
                                i.text = text
                            if '{ContractMonthKAZ}' in i.text:
                                text = i.text.replace('{ContractMonthKAZ}', f'{self.translate_text(contract.ContractDate.strftime("%B"), "kk")}')
                                i.text = text
                            if '{ContractMonthENG}' in i.text:
                                text = i.text.replace('{ContractMonthENG}', f'{self.translate_text(contract.ContractDate.strftime("%B"), "en")}')
                                i.text = text
                            if '{EduYear}' in i.text:
                                text = i.text.replace('{EduYear}', f'{contract.EduYearID.sEduYear}')
                                i.text = text
                            if '{ParentFullName}' in i.text:
                                text = i.text.replace('{ParentFullName}', f'{parent.full_name}')
                                i.text = text
                            if '{StudentFullName}' in i.text:
                                text = i.text.replace('{StudentFullName}', f'{student.full_name}')
                                i.text = text
                            if '{ContractDopAmount}' in i.text:
                                text = i.text.replace('{ContractDopAmount}', f'{int(contract_dop_amount)}')
                                i.text = text
                            if '{ContractDopAmountWords}' in i.text:
                                text = i.text.replace('{ContractDopAmountWords}',
                                                      f'{num2words(int(contract_dop_amount), lang="ru")}')
                                i.text = text
                            if '{ContractDopAmountWordsKaz}' in i.text:
                                text = i.text.replace('{ContractDopAmountWordsKaz}',
                                                      f'{num2words(int(contract_dop_amount), lang="kz")}')
                                i.text = text
                            if '{ContractAmount}' in i.text:
                                text = i.text.replace('{ContractAmount}', f'{int(contract.ContractAmount)}')
                                i.text = text
                            if '{ContractAmountWords}' in i.text:
                                text = i.text.replace('{ContractAmountWords}',
                                                      f'{num2words(int(contract.ContractAmount), lang="ru")}')
                                i.text = text
                            if '{ContractAmountWordsKaz}' in i.text:
                                text = i.text.replace('{ContractAmountWordsKaz}',
                                                      f'{num2words(int(contract.ContractAmount), lang="kz")}')
                                i.text = text
                            if '{ContractAmountWordsEng}' in i.text:
                                text = i.text.replace('{ContractAmountWordsEng}',
                                                      f'{num2words(int(contract.ContractAmount), lang="en")}')
                                i.text = text
                            if '{ContractSum}' in i.text:
                                text = i.text.replace('{ContractSum}', f'{int(contract.ContractSum)}')
                                i.text = text
                            if '{ContractSumWords}' in i.text:
                                text = i.text.replace('{ContractSumWords}',
                                                      f'{num2words(int(contract.ContractSum), lang="ru")}')
                                i.text = text
                            if '{ContractSumWordsKaz}' in i.text:
                                text = i.text.replace('{ContractSumWordsKaz}',
                                                      f'{num2words(int(contract.ContractSum), lang="kz")}')
                                i.text = text
                            if '{ContractSumWordsEng}' in i.text:
                                text = i.text.replace('{ContractSumWordsEng}',
                                                      f'{num2words(int(contract.ContractSum), lang="en")}')
                                i.text = text
                            if '{ContractAmountWithDiscount}' in i.text:
                                text = i.text.replace('{ContractAmountWithDiscount}', f'{int(contract_amount_with_discount)}')
                                i.text = text
                            if '{ContractAmountWithDiscountWords}' in i.text:
                                text = i.text.replace('{ContractAmountWithDiscountWords}',
                                                      f'{num2words(int(contract_amount_with_discount), lang="ru")}')
                                i.text = text
                            if '{ContractAmountWithDiscountWordsKaz}' in i.text:
                                text = i.text.replace('{ContractAmountWithDiscountWordsKaz}',
                                                      f'{num2words(int(contract_amount_with_discount), lang="kz")}')
                                i.text = text
                            if '{ContractAmountWithDiscountWordsEng}' in i.text:
                                text = i.text.replace('{ContractAmountWithDiscountWordsEng}',
                                                      f'{num2words(int(contract_amount_with_discount), lang="en")}')
                                i.text = text
                            if '{ContractContr}' in i.text:
                                if contract.ContSum is None:
                                    contract.ContSum = 0
                                text = i.text.replace('{ContractContr}', f'{int(contract.ContSum)}')
                                i.text = text
                            if '{ContractContrWords}' in i.text:
                                if contract.ContSum is None:
                                    contract.ContSum = 0
                                text = i.text.replace('{ContractContrWords}',
                                                      f'{num2words(int(contract.ContSum), lang="ru")}')
                                i.text = text
                            if '{ContractContrWordsKaz}' in i.text:
                                if contract.ContSum is None:
                                    contract.ContSum = 0
                                text = i.text.replace('{ContractContrWordsKaz}',
                                                      f'{num2words(int(contract.ContSum), lang="kz")}')
                                i.text = text
                            if '{ContractContrWordsEng}' in i.text:
                                if contract.ContSum is None:
                                    contract.ContSum = 0
                                text = i.text.replace('{ContractContrWordsEng}',
                                                      f'{num2words(int(contract.ContSum), lang="en")}')
                                i.text = text
                            if '{StudentIIN}' in i.text:
                                text = i.text.replace('{StudentIIN}', f'{str(student.iin)}')
                                i.text = text
                            if '{StudentAddress}' in i.text:
                                text = i.text.replace('{StudentAddress}', f'{parent.address}')
                                i.text = text
                            if '{StudentPhoneNumber}' in i.text:
                                if student.phone is not None:
                                    text = i.text.replace('{StudentPhoneNumber}', f'{student.phone}')
                                else:
                                    text = i.text.replace('{StudentPhoneNumber}', '-')
                                i.text = text
                            if '{ParentAddress}' in i.text:
                                text = i.text.replace('{ParentAddress}', f'{parent.address}')
                                i.text = text
                            if '{ParentPhoneNumber}' in i.text:
                                text = i.text.replace('{ParentPhoneNumber}', f'{parent.phone}')
                                i.text = text
                            if '{ParentIIN}' in i.text:
                                text = i.text.replace('{ParentIIN}', f'{parent.iin}')
                                i.text = text
                            if '{ParentPassport}' in i.text:
                                text = i.text.replace('{ParentPassport}', ParentPassport)
                                i.text = text
                            if '{ParentPassportKAZ}' in i.text:
                                text = i.text.replace('{ParentPassportKAZ}', ParentPassportKAZ)
                                i.text = text
                            if '{ParentPassportENG}' in i.text:
                                text = i.text.replace('{ParentPassportENG}', ParentPassportENG)
                                i.text = text

        self.change_docx_contract_month_pay(doc, contract, is_dop_contract)

    def change_content(self, request, contract_num, contract, student, parent, is_dop_contract):
        contract_file = GetContractFromDBService.get_contract(contract_num, is_dop_contract)

        if contract_file is not None:
            pdf_file = contract_file.file

        else:
            contract_payment_type = contract.PaymentTypeID.sPaymentType

            if contract.SchoolID.sSchool_language == 'Казахское отделение':
                if is_dop_contract:
                    if contract_payment_type == 'Оплата по месячно':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_по_месячно.docx'
                    elif contract_payment_type == 'Оплата по квартально':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_по_квартально.docx'
                    else:
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_за_год.docx'
                else:
                    if contract_payment_type == 'Оплата по месячно':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_по_месячно.docx'
                    elif contract_payment_type == 'Оплата по квартально':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_по_квартально.docx'
                    else:
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_за_год.docx'
            else:
                if contract.SchoolID.sSchool_direct == 'Кембридж':
                    if contract_payment_type == 'Оплата по месячно':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Кэмбридж_2023_2024_оплата_по_месячно.docx'
                    elif contract_payment_type == 'Оплата по квартально':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Кэмбридж_2023_2024_оплата_по_квартально.docx'
                    else:
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Кэмбридж_2023_2024_оплата_за_год.docx'
                elif is_dop_contract:
                    if contract.SchoolID.sSchool_direct == 'Лингвинистический':
                        if contract_payment_type == 'Оплата по месячно':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_по_месячно.docx'
                        elif contract_payment_type == 'Оплата по квартально':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_по_квартально.docx'
                        else:
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_за_год.docx'
                    elif contract.SchoolID.sSchool_direct == 'Физико-математический' or contract.SchoolID.sSchool_direct == 'Физико-Математическая':
                        if contract_payment_type == 'Оплата по месячно':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Физмат_по_месячно.docx'
                        elif contract_payment_type == 'Оплата по квартально':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Физмат_по_квартально.docx'
                        else:
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_Физмат_за_год.docx'
                    elif contract.SchoolID.sSchool_direct == 'Американская школа Advanced Placement':
                        if contract_payment_type == 'Оплата по месячно':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_по_месячно.docx'
                        elif contract_payment_type == 'Оплата по квартально':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_по_квартально.docx'
                        else:
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_за_год.docx'
                    elif contract.SchoolID.sSchool_direct == 'IT-школа на Кекилбайулы':
                        if contract_payment_type == 'Оплата по месячно':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_по_месячно.docx'
                        elif contract_payment_type == 'Оплата по квартально':
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_по_квартально.docx'
                        else:
                            docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_за_год.docx'
                    else:
                        print('Шаблон договора не найден!')
                        return Response({"message": "Шаблон договора не найден!"}, status=status.HTTP_403_FORBIDDEN)
                else:
                    if contract_payment_type == 'Оплата по месячно':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_месячно.docx'
                    elif contract_payment_type == 'Оплата по квартально':
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_квартально.docx'
                    else:
                        docx_file = 'apps/contract/templates/contract/Шаблон_Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_за_год.docx'

            doc = Document(docx_file)

            if doc is not None:
                try:
                    print("before change docx document")
                    self.change_docx_document(doc, contract, student, parent, is_dop_contract)
                    print("after change docx document")
                except ValueError:
                    print('Ошибка при изменении содержимого документа!')
                    return Response({'error': 'Ошибка при изменении содержимого документа!'}, status=status.HTTP_403_FORBIDDEN)
            else:
                print('Не найден шаблон договора!')
                return Response({'error': 'Не найден шаблон договора!'}, status=status.HTTP_403_FORBIDDEN)

            docx_output_path = f'contracts/version/docx/contract_{contract_num}.docx'

            pdf_directory = "contracts/version/pdf"
            pdf_output_path = f'{pdf_directory}/contract_{contract_num}.pdf'

            pdf_file = self.save_document(request, contract_num, doc, docx_output_path, pdf_directory, pdf_output_path, is_dop_contract)

        return pdf_file


class ContractDownloadService:
    """
        Сервис для скачивания договора студента.
        Договор ищется по номеру договора.
        Файл скачивается в формате PDF.
    """

    def __init__(self, contract_student) -> None:
        self.contract_student = contract_student

    @staticmethod
    def check_exist_student(contract):
        try:
            student_id = contract.StudentID
            student = StudentMS.objects.using('ms_sql').get(id=student_id.id)
        except AttributeError:
            raise Exception('Студент не найден!')
        except ObjectDoesNotExist:
            raise Exception('Студент не найден!')

        return student

    def check_exists_parent(self, contract):
        try:
            student = self.check_exist_student(contract)
            parent_id = ParentMS.objects.using('ms_sql').get(id=student.parent_id.id)
            parent = ParentMS.objects.using('ms_sql').get(id=parent_id.id)
        except AttributeError:
            raise Exception('Родитель не найден!')
        except ObjectDoesNotExist:
            raise Exception('Родитель не найден!')

        return parent

    def contract_download(self, request, contract_num, is_dop_contract):
        contract = self.contract_student

        if contract is None:
            return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)

        if contract.ContractDate.year < datetime.now().year - 1:
            return Response({'error': 'Скачать договор можно только на текущий год!'}, status=status.HTTP_403_FORBIDDEN)

        try:
            student = self.check_exist_student(contract)
            parent = self.check_exists_parent(contract)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_403_FORBIDDEN)

        change_doc_service = ChangeDocumentContentService()

        if is_dop_contract:
            contract_num = contract_num.replace('/', '-')
            contract_file = ContractDopFileUser.objects.filter(contractNum=contract_num).last()
            if contract_file:
                file = contract_file.file
            else:
                change_doc_service.change_content(request, contract_num, contract, student, parent, is_dop_contract=True)

                contract_file = ContractDopFileUser.objects.filter(contractNum=contract_num).last()
                if contract_file:
                    file = contract_file.file
                else:
                    print("Договор не найден!")
                    return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)
        else:
            contract_file = ContractFileUser.objects.filter(contractNum=contract_num).last()
            if contract_file:
                file = contract_file.file
            else:
                change_doc_service.change_content(request, contract_num, contract, student, parent, is_dop_contract=False)

                contract_file = ContractFileUser.objects.filter(contractNum=contract_num).last()
                if contract_file:
                    file = contract_file.file
                else:
                    print("Договор не найден!")
                    return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)

        response = FileResponse(file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{contract_num}.pdf"'
        return response

    def generate_contract_with_qr_code(self, request, contract_num, qr_code, qr_code_director_omarov, qr_code_director_serikov, is_dop_contract):
        """
            Получаем PDF файл из базы данных, затем конвертируем его в docx.
            Изменяем содержимое документа и опять конвертируем в PDF.
            Обновляем файл в базе данных.
        """

        if is_dop_contract:
            try:
                contract = self.contract_student
            except ContractDopMS.DoesNotExist:
                return Response({'error': 'Дополнительный договор не найден!'}, status=status.HTTP_403_FORBIDDEN)
        else:
            try:
                contract = self.contract_student
            except ObjectDoesNotExist:
                return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)

        try:
            student = self.check_exist_student(contract)
            parent = self.check_exists_parent(contract)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_403_FORBIDDEN)

        try:
            contract_payment_type = contract.PaymentTypeID.sPaymentType
        except AttributeError:
            raise ValueError('Не найден тип оплаты!')
        try:
            contract_school_direct = contract.SchoolID.sSchool_direct
        except AttributeError:
            raise ValueError('Не найдено направление школы!')
        try:
            contract_school_language = contract.SchoolID.sSchool_language
        except AttributeError:
            raise ValueError('Не найден язык школы!')
        if contract_payment_type == 'Оплата по месячно':
            if contract_school_language == 'Казахское отделение':
                if is_dop_contract:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_ТОО_по_месячно.docx'
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_по_месячно.docx'
            else:
                if contract_school_direct == 'Кембридж':
                    docx_file = 'apps/contract/templates/contract/Договор оказания образовательных услуг Кэмбридж 2025-2026 УО_по_месячно.docx'
                elif is_dop_contract:
                    if contract_school_direct == 'Лингвинистический':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_по_месячно.docx'
                    elif contract_school_direct == 'Физико-математический' or contract_school_direct == 'Физико-Математическая':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Физмат_Нур_по_месячно.docx'
                    elif contract_school_direct == 'Американская школа Advanced Placement':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_по_месячно.docx'
                    elif contract_school_direct == 'IT-школа на Кекилбайулы':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_по_месячно.docx'
                    else:
                        return Response({"message": "Шаблон договора не найден!"}, status=status.HTTP_403_FORBIDDEN)
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_месячно.docx'
        elif contract_payment_type == 'Оплата по квартально':
            if contract_school_language == 'Казахское отделение':
                if is_dop_contract:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_ТОО_по_квартально.docx'
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_по_квартально.docx'
            else:
                if contract_school_direct == 'Кембридж':
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_Кэмбридж_2023_2024_оплата_по_квартально.docx'
                elif is_dop_contract:
                    if contract_school_direct == 'Лингвинистический':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_по_квартально.docx'
                    elif contract_school_direct == 'Физико-математический' or contract_school_direct == 'Физико-Математическая':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Физмат_Нур_по_квартально.docx'
                    elif contract_school_direct == 'Американская школа Advanced Placement':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_по_квартально.docx'
                    elif contract_school_direct == 'IT-школа на Кекилбайулы':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_по_квартально.docx'
                    else:
                        return Response({"message": "Шаблон договора не найден!"}, status=status.HTTP_403_FORBIDDEN)
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_по_квартально.docx'
        else:
            if contract_school_language == 'Казахское отделение':
                if is_dop_contract:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_КАЗ_ОТД_ТОО_за_год.docx'
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_КАЗ_ОТД_ТОО_за_год.docx'
            else:
                if contract_school_direct == 'Кембридж':
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_Кэмбридж_2023_2024_оплата_за_год.docx'
                elif is_dop_contract:
                    if contract_school_direct == 'Лингвинистический':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Лингво_2023_за_год.docx'
                    elif contract_school_direct == 'Физико-математический' or contract_school_direct == 'Физико-Математическая':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_Физмат_Нур_за_год.docx'
                    elif contract_school_direct == 'Американская школа Advanced Placement':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_AP_2023_2024_за_год.docx'
                    elif contract_school_direct == 'IT-школа на Кекилбайулы':
                        docx_file = 'apps/contract/templates/contract/Договор_оказания_дополнительных_образовательных_услуг_IT_отделение_за_год.docx'
                    else:
                        return Response({"message": "Шаблон договора не найден!"}, status=status.HTTP_403_FORBIDDEN)
                else:
                    docx_file = 'apps/contract/templates/contract/Договор_оказания_образовательных_услуг_Школа_2023_2024_оплата_за_год.docx'

        doc = Document(docx_file)

        ChangeDocumentContentService().change_docx_document(doc, contract, student, parent, is_dop_contract)

        for paragraph in doc.paragraphs:
            inline = paragraph
            for i in inline.runs:
                if '{QRCode}' in i.text:
                    i.text = ''
                    image_stream = BytesIO(qr_code)
                    i.add_picture(image_stream, width=Inches(2.0), height=Inches(2.0))
                if '{QRCodeTextRus}' in i.text:
                    text = 'QR-код содержит данные об электронно-цифровой подписи подписанта'
                    i.text = text
                if '{QRCodeTextKaz}' in i.text:
                    text = 'QR-кодта қол қоюшының электрондық-цифрлық қолтаңбасы туралы деректер қамтылады'
                    i.text = text
                if 'QRCodeDataSigned' in i.text:
                    from apps.contract.services_eds import SignContractWithEDSService

                    signed_contract_service = SignContractWithEDSService(self.contract_student)
                    qr_code_data_signed = signed_contract_service.generate_qr_code_data_signed(contract_num=contract_num)
                    qr_code_image = BytesIO(qr_code_data_signed)
                    i.add_picture(qr_code_image, width=Inches(1.3), height=Inches(1.3))
                if '{police_kaz}' in i.text:
                    text = 'Осы құжат «Электрондық құжат және электрондық цифрлық қолтаңба туралы» Қазақстан Республикасының 2003 жылғы 7 қаңтардағы N 370-II Заңы 7 бабының 1 тармағына сәйкес қағаз тасығыштағы құжатпен бірдей.'
                    i.text = text
                if '{police_rus}' in i.text:
                    text = 'Данный документ согласно пункту 1 статьи 7 ЗРК от 7 января 2003 года «Об электронном документе и электронной цифровой подписи» равнозначен документу на бумажном носителе.'
                    i.text = text
                if '{QRcodeDirector}' in i.text:
                    i.text = ''
                    image_stream = BytesIO(qr_code_director_omarov)
                    i.add_picture(image_stream, width=Inches(2.0), height=Inches(2.0))
                if '{QRCodeDirector2}' in i.text:
                    i.text = ''
                    image_stream = BytesIO(qr_code_director_serikov)
                    i.add_picture(image_stream, width=Inches(2.0), height=Inches(2.0))

        doc.save(f'contracts/version/docx/{contract_num}.docx')
        ChangeDocumentContentService.docx_to_pdf(f'contracts/version/docx/{contract_num}.docx', 'contracts/version/pdf')

        os.remove(f'contracts/version/docx/{contract_num}.docx')

        pdf_file = f'contracts/version/pdf/{contract_num}.pdf'

        if is_dop_contract:
            with open(pdf_file, 'rb') as pdf:
                if not ContractDopFileUser.objects.filter(contractNum=contract_num).exists():
                    contract_file = ContractDopFileUser.objects.create(
                        user=request.user,
                        contractNum=contract_num,
                        file=ContentFile(pdf.read(), name=f'{contract_num}.pdf')
                    )
                    contract_file.save()

                else:
                    contract_file = ContractDopFileUser.objects.filter(contractNum=contract_num).last()
                    if contract_file:
                        contract_file.date = datetime.now()
                        contract_file.file = ContentFile(pdf.read(), name=f'{contract_num}.pdf')
                        contract_file.save()
                    else:
                        return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)
        else:
            with open(pdf_file, 'rb') as pdf:
                if not ContractFileUser.objects.filter(contractNum=contract_num).exists():
                    contract_file = ContractFileUser.objects.create(
                        user=request.user,
                        contractNum=contract_num,
                        file=ContentFile(pdf.read(), name=f'{contract_num}.pdf')
                    )
                    contract_file.save()

                else:
                    contract_file = ContractFileUser.objects.filter(contractNum=contract_num).last()
                    if contract_file:
                        contract_file.date = datetime.now()
                        contract_file.file = ContentFile(pdf.read(), name=f'{contract_num}.pdf')
                        contract_file.save()
                    else:
                        return Response({'error': 'Договор не найден!'}, status=status.HTTP_403_FORBIDDEN)

        os.remove(pdf_file)

        return JsonResponse({'message': 'Договор успешно сгенерирован!'}, status=status.HTTP_200_OK)
